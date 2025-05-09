from flask import render_template, request, send_file
from flask_socketio import emit
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO

from . import editor_bp

# Carregar dados do Excel
try:
    fornecedores_df = pd.read_excel("Uploads/fornecedores.xlsx")
    empresas = fornecedores_df['EMPRESA'].unique().tolist()
    produtos = fornecedores_df['PRODUTO'].unique().tolist()
    marcas = fornecedores_df['MARCA'].unique().tolist()
except Exception as e:
    print(f"Erro ao carregar dados do Excel: {e}")
    fornecedores_df = pd.DataFrame(columns=["EMPRESA", "PRODUTO", "MARCA"])
    empresas = []
    produtos = []
    marcas = []

@editor_bp.route('/editor', methods=['GET', 'POST'])
def editor():
    if request.method == 'POST':
        format = request.form.get('format')
        content = request.form.get('content')
        empresa = request.form.get('empresa')
        produto = request.form.get('produto')
        marca = request.form.get('marca')
        
        # Gerar arquivo baseado no formato solicitado
        if format == 'pdf':
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            c.drawString(100, 750, f"Empresa: {empresa}")
            c.drawString(100, 730, f"Produto: {produto}")
            c.drawString(100, 710, f"Marca: {marca}")
            c.drawString(100, 690, "Conteúdo:")
            y = 670
            for line in content.split('\n'):
                c.drawString(100, y, line)
                y -= 20
            c.save()
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name="laudo.pdf", mimetype="application/pdf")
        
        elif format == 'word':
            doc = Document()
            doc.add_paragraph(f"Empresa: {empresa}")
            doc.add_paragraph(f"Produto: {produto}")
            doc.add_paragraph(f"Marca: {marca}")
            doc.add_paragraph("Conteúdo:")
            doc.add_paragraph(content)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name="laudo.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        elif format == 'excel':
            df = pd.DataFrame({
                "Empresa": [empresa],
                "Produto": [produto],
                "Marca": [marca],
                "Conteúdo": [content]
            })
            buffer = BytesIO()
            df.to_excel(buffer, index=False)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name="laudo.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    return render_template('editor.html', 
                         empresas=empresas, 
                         produtos=produtos, 
                         marcas=marcas, 
                         fornecedores_df=fornecedores_df)

from app import socketio

@socketio.on('update_content')
def handle_update_content(data):
    emit('content_updated', data, broadcast=True)