"""
Módulo para Edição Online de Documentos

Este módulo contém funções para processar e manipular documentos Excel, Word e PDF,
extraindo seus campos e aplicando valores editados.
"""

import os
import re
import io
import json
import base64
import tempfile
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

# Importações para manipulação de documentos
import docx
import PyPDF2
import pandas as pd
from flask import current_app
import openpyxl
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def extract_data_from_excel(file_path: str) -> Dict[str, Any]:
    """
    Extrai dados de um arquivo Excel, incluindo folhas, células e campos.
    
    Args:
        file_path: Caminho para o arquivo Excel
        
    Returns:
        Dicionário com os dados extraídos
    """
    # Abrir o arquivo Excel
    workbook = openpyxl.load_workbook(file_path)
    
    # Obter nome da aba ativa
    active_sheet_name = workbook.active.title
    
    # Preparar dicionário para os dados de todas as abas
    sheets_data = {}
    fields = []
    
    # Para cada aba do arquivo
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        
        # Preparar dados da aba
        sheet_data = {
            'columns': [],
            'data': []
        }
        
        # Obter dimensões da aba
        max_row = sheet.max_row
        max_col = sheet.max_column
        
        # Limitar a quantidade de células para melhor desempenho
        max_row = min(max_row, 100)
        max_col = min(max_col, 50)
        
        # Preencher colunas (A, B, C, ...)
        for col in range(1, max_col + 1):
            sheet_data['columns'].append(col)
        
        # Para cada linha da aba
        for row in range(1, max_row + 1):
            row_data = []
            
            # Para cada coluna da aba
            for col in range(1, max_col + 1):
                cell = sheet.cell(row=row, column=col)
                cell_value = cell.value if cell.value is not None else ""
                
                # Verificar se a célula é um campo preenchível
                is_field = False
                if isinstance(cell_value, str) and (
                    cell_value.startswith("{{") and cell_value.endswith("}}") or
                    cell_value.startswith("[") and cell_value.endswith("]") or
                    "_______" in cell_value or
                    "______" in cell_value or
                    "__________" in cell_value
                ):
                    is_field = True
                    # Extrair nome do campo
                    field_name = cell_value
                    if field_name.startswith("{{") and field_name.endswith("}}"):
                        field_name = field_name[2:-2].strip()
                    elif field_name.startswith("[") and field_name.endswith("]"):
                        field_name = field_name[1:-1].strip()
                    
                    # Limpar nome do campo
                    field_name = field_name.replace("_", "").strip()
                    if not field_name:
                        field_name = f"Campo {row}{chr(64+col)}"
                    
                    # Criar ID único para o campo
                    field_id = f"excel_{sheet_name}_{row}_{col}"
                    
                    # Adicionar à lista de campos
                    fields.append({
                        'id': field_id,
                        'name': field_name,
                        'sheet': sheet_name,
                        'row': row,
                        'col': col,
                        'value': ""
                    })
                
                # Adicionar dados da célula
                row_data.append({
                    'value': str(cell_value),
                    'is_field': is_field,
                    'row': row,
                    'col': col
                })
            
            # Adicionar linha aos dados da aba
            sheet_data['data'].append(row_data)
        
        # Adicionar dados da aba ao dicionário de abas
        sheets_data[sheet_name] = sheet_data
    
    # Retornar dados extraídos
    return {
        'active_sheet': active_sheet_name,
        'sheets': sheets_data,
        'fields': fields
    }

def apply_excel_fields(input_path: str, output_path: str, fields: List[Dict[str, Any]]) -> None:
    """
    Aplica valores editados a um arquivo Excel.
    
    Args:
        input_path: Caminho para o arquivo Excel original
        output_path: Caminho para salvar o arquivo Excel editado
        fields: Lista de campos com valores editados
    """
    # Abrir o arquivo Excel
    workbook = openpyxl.load_workbook(input_path)
    
    # Para cada campo editado
    for field in fields:
        # Verificar se o campo é do Excel
        if not field['id'].startswith('excel_'):
            continue
        
        # Extrair informações do campo
        parts = field['id'].split('_')
        if len(parts) >= 4:
            # Formato: excel_sheet_row_col
            sheet_name = parts[1]
            row = int(parts[2])
            col = int(parts[3])
            
            # Verificar se a aba existe
            if sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Aplicar valor editado
                edited_value = field.get('edited_value', "")
                sheet.cell(row=row, column=col).value = edited_value
    
    # Salvar o arquivo Excel editado
    workbook.save(output_path)

def extract_data_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extrai dados de um arquivo Word, incluindo parágrafos, tabelas e campos.
    
    Args:
        file_path: Caminho para o arquivo Word
        
    Returns:
        Dicionário com os dados extraídos
    """
    # Abrir o arquivo Word
    doc = docx.Document(file_path)
    
    # Preparar dados
    paragraphs_data = []
    tables_data = []
    fields = []
    
    # Padrões para detectar campos preenchíveis
    field_patterns = [
        r"\{\{([^}]+)\}\}",  # {{campo}}
        r"\[([^]]+)\]",      # [campo]
        r"_{5,}",            # ______
        r"__________"        # __________
    ]
    
    # Extrair parágrafos
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Verificar se o parágrafo contém um campo preenchível
        is_field = False
        for pattern in field_patterns:
            if re.search(pattern, text):
                is_field = True
                
                # Extrair nome do campo
                field_name = text
                match = re.search(r"\{\{([^}]+)\}\}", text)
                if match:
                    field_name = match.group(1).strip()
                else:
                    match = re.search(r"\[([^]]+)\]", text)
                    if match:
                        field_name = match.group(1).strip()
                    else:
                        field_name = text.replace("_", "").strip()
                        if not field_name:
                            field_name = f"Parágrafo {i+1}"
                
                # Criar ID único para o campo
                field_id = f"docx_para_{i}"
                
                # Adicionar à lista de campos
                fields.append({
                    'id': field_id,
                    'name': field_name,
                    'paragraph_index': i,
                    'value': ""
                })
        
        # Adicionar dados do parágrafo
        paragraphs_data.append({
            'text': text,
            'is_field': is_field,
            'paragraph_index': i
        })
    
    # Extrair tabelas
    for table_index, table in enumerate(doc.tables):
        table_data = []
        
        # Para cada linha da tabela
        for row_index, row in enumerate(table.rows):
            row_data = []
            
            # Para cada célula da linha
            for col_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                
                # Verificar se a célula contém um campo preenchível
                is_field = False
                for pattern in field_patterns:
                    if re.search(pattern, text):
                        is_field = True
                        
                        # Extrair nome do campo
                        field_name = text
                        match = re.search(r"\{\{([^}]+)\}\}", text)
                        if match:
                            field_name = match.group(1).strip()
                        else:
                            match = re.search(r"\[([^]]+)\]", text)
                            if match:
                                field_name = match.group(1).strip()
                            else:
                                field_name = text.replace("_", "").strip()
                                if not field_name:
                                    field_name = f"Tabela {table_index+1} Célula {row_index+1}x{col_index+1}"
                        
                        # Criar ID único para o campo
                        field_id = f"docx_table_{table_index}_cell_{row_index}_{col_index}"
                        
                        # Adicionar à lista de campos
                        fields.append({
                            'id': field_id,
                            'name': field_name,
                            'table_index': table_index,
                            'row_index': row_index,
                            'col_index': col_index,
                            'value': ""
                        })
                
                # Adicionar dados da célula
                row_data.append({
                    'text': text,
                    'is_field': is_field,
                    'row_index': row_index,
                    'col_index': col_index
                })
            
            # Adicionar linha aos dados da tabela
            table_data.append(row_data)
        
        # Adicionar dados da tabela
        tables_data.append({
            'table_index': table_index,
            'data': table_data
        })
    
    # Retornar dados extraídos
    return {
        'paragraphs': paragraphs_data,
        'tables': tables_data,
        'fields': fields
    }

def apply_docx_fields(input_path: str, output_path: str, fields: List[Dict[str, Any]]) -> None:
    """
    Aplica valores editados a um arquivo Word.
    
    Args:
        input_path: Caminho para o arquivo Word original
        output_path: Caminho para salvar o arquivo Word editado
        fields: Lista de campos com valores editados
    """
    # Abrir o arquivo Word
    doc = docx.Document(input_path)
    
    # Mapear campos por tipo e índice
    paragraph_fields = {}
    table_fields = {}
    
    for field in fields:
        if field['id'].startswith('docx_para_'):
            paragraph_index = int(field['id'].split('_')[2])
            paragraph_fields[paragraph_index] = field.get('edited_value', "")
        
        elif field['id'].startswith('docx_table_'):
            # Formato: docx_table_table_index_cell_row_index_col_index
            parts = field['id'].split('_')
            table_index = int(parts[2])
            row_index = int(parts[5])
            col_index = int(parts[6])
            
            if table_index not in table_fields:
                table_fields[table_index] = {}
            
            if row_index not in table_fields[table_index]:
                table_fields[table_index][row_index] = {}
            
            table_fields[table_index][row_index][col_index] = field.get('edited_value', "")
    
    # Aplicar valores aos parágrafos
    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        if paragraph_index in paragraph_fields:
            # Substituir conteúdo do parágrafo
            paragraph.text = paragraph_fields[paragraph_index]
    
    # Aplicar valores às tabelas
    for table_index, table in enumerate(doc.tables):
        if table_index in table_fields:
            for row_index, row in enumerate(table.rows):
                if row_index in table_fields[table_index]:
                    for col_index, cell in enumerate(row.cells):
                        if col_index in table_fields[table_index][row_index]:
                            # Substituir conteúdo da célula
                            cell.text = table_fields[table_index][row_index][col_index]
    
    # Salvar o arquivo Word editado
    doc.save(output_path)

def extract_data_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extrai dados de um arquivo PDF, incluindo campos preenchíveis.
    
    Args:
        file_path: Caminho para o arquivo PDF
        
    Returns:
        Dicionário com os dados extraídos
    """
    # Abrir o arquivo PDF
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        
        # Extrair campos de formulário, se existirem
        fields = []
        
        # Padrões para detectar campos em textos
        field_patterns = [
            r"\{\{([^}]+)\}\}",  # {{campo}}
            r"\[([^]]+)\]",      # [campo]
            r"_{5,}",            # ______
            r"__________"        # __________
        ]
        
        # Tentar encontrar campos de formulário interativo
        if '/AcroForm' in reader.trailer['/Root']:
            # PDF tem campos de formulário
            if '/Fields' in reader.trailer['/Root']['/AcroForm']:
                form_fields = reader.trailer['/Root']['/AcroForm']['/Fields']
                
                for i, field_ref in enumerate(form_fields):
                    field_dict = field_ref.get_object()
                    field_type = field_dict.get('/FT', '')
                    field_name = field_dict.get('/T', f'Campo {i+1}')
                    
                    # Criar ID único para o campo
                    field_id = f"pdf_form_{i}"
                    
                    # Adicionar à lista de campos
                    fields.append({
                        'id': field_id,
                        'name': field_name,
                        'type': field_type,
                        'page': 0,  # PDF interativo, página não é relevante
                        'value': ""
                    })
        
        # Se não houver campos interativos, tentar extrair texto e encontrar padrões de campos
        if not fields:
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                
                # Dividir o texto em linhas
                lines = text.split('\n')
                
                for i, line in enumerate(lines):
                    # Verificar se a linha contém um padrão de campo
                    for pattern in field_patterns:
                        matches = re.finditer(pattern, line)
                        for match in matches:
                            # Extrair nome do campo
                            matched_text = match.group(0)
                            field_name = matched_text
                            
                            if matched_text.startswith("{{") and matched_text.endswith("}}"):
                                field_name = matched_text[2:-2].strip()
                            elif matched_text.startswith("[") and matched_text.endswith("]"):
                                field_name = matched_text[1:-1].strip()
                            else:
                                field_name = matched_text.replace("_", "").strip()
                                if not field_name:
                                    field_name = f"Campo {page_num+1}_{i+1}"
                            
                            # Criar ID único para o campo
                            field_id = f"pdf_text_{page_num}_{i}_{match.start()}"
                            
                            # Adicionar à lista de campos
                            fields.append({
                                'id': field_id,
                                'name': field_name,
                                'page': page_num,
                                'line': i,
                                'position': match.start(),
                                'matched_text': matched_text,
                                'value': ""
                            })
        
        # Retornar dados extraídos
        return {
            'fields': fields
        }

def apply_pdf_fields(input_path: str, output_path: str, fields: List[Dict[str, Any]]) -> None:
    """
    Aplica valores editados a um arquivo PDF.
    
    Nota: Esta função cria uma anotação simples sobre o PDF original com os valores editados,
    pois a edição direta de PDFs é complexa e requer bibliotecas mais avançadas.
    
    Args:
        input_path: Caminho para o arquivo PDF original
        output_path: Caminho para salvar o arquivo PDF editado
        fields: Lista de campos com valores editados
    """
    # Para PDFs, criamos uma anotação simples sobre o PDF original
    # Esta é uma solução temporária, pois editar PDFs diretamente é complexo
    
    # Filtrar apenas campos com valores editados
    edited_fields = [f for f in fields if f.get('edited_value')]
    
    if not edited_fields:
        # Se não houver campos editados, apenas copiar o arquivo original
        with open(input_path, 'rb') as infile, open(output_path, 'wb') as outfile:
            outfile.write(infile.read())
        return
    
    # Método simplificado: criar um PDF com anotações e depois mesclá-lo com o original
    # Em uma implementação mais robusta, usaríamos uma biblioteca específica para formulários PDF
    
    # Criar um PDF temporário com as anotações dos campos preenchidos
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
        temp_path = temp_file.name
    
    # Criar um PDF com anotações
    c = canvas.Canvas(temp_path, pagesize=letter)
    c.setFont("Helvetica", 10)
    
    # Adicionar informações dos campos editados
    y = 750  # Posição inicial
    c.drawString(50, y, f"Formulário preenchido em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    y -= 20
    c.drawString(50, y, "Campos preenchidos:")
    y -= 20
    
    for field in edited_fields:
        field_text = f"{field.get('name', 'Campo')}: {field.get('edited_value', '')}"
        c.drawString(70, y, field_text)
        y -= 15
        
        # Se chegamos ao final da página, criar uma nova
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = 750
    
    c.save()
    
    # Mesclar o PDF original com o PDF de anotações
    with open(input_path, 'rb') as input_file, open(temp_path, 'rb') as temp_file, open(output_path, 'wb') as output_file:
        original_pdf = PyPDF2.PdfReader(input_file)
        annotation_pdf = PyPDF2.PdfReader(temp_file)
        
        output_pdf = PyPDF2.PdfWriter()
        
        # Adicionar todas as páginas do PDF original
        for page in original_pdf.pages:
            output_pdf.add_page(page)
        
        # Adicionar páginas de anotação
        for page in annotation_pdf.pages:
            output_pdf.add_page(page)
        
        # Salvar o PDF mesclado
        output_pdf.write(output_file)
    
    # Remover arquivo temporário
    os.unlink(temp_path)

# Funções auxiliares
def get_file_size_display(size_in_bytes: int) -> str:
    """
    Formata o tamanho do arquivo para exibição.
    
    Args:
        size_in_bytes: Tamanho em bytes
        
    Returns:
        String formatada com o tamanho do arquivo
    """
    if size_in_bytes < 1024:
        return f"{size_in_bytes} bytes"
    elif size_in_bytes < 1024 * 1024:
        return f"{size_in_bytes / 1024:.1f} KB"
    else:
        return f"{size_in_bytes / (1024 * 1024):.1f} MB"