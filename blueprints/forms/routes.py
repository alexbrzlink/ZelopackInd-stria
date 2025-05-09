import logging
logger = logging.getLogger(__name__)

import os
import mimetypes
import json
import tempfile
import shutil
import datetime
from flask import render_template, send_file, abort, Response, jsonify, request, redirect, url_for, flash
from flask_login import login_required, current_user
from flask_wtf.csrf import generate_csrf
from . import forms_bp
from app import db, csrf
from models import StandardFields, FormPreset
from models import FormPreset
import openpyxl
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import docx

# Diretório base dos formulários
FORMS_DIR = os.path.join(os.getcwd(), 'extracted_forms')

@forms_bp.route('/')
@login_required
def index():
    """Página principal para visualização de formulários de ordem de produção."""
    # Esta página agora irá mostrar apenas os formulários de ordem de produção
    # Os outros formulários foram movidos para a aba Documentos
    
    ordem_producao_dir = None
    for item in os.listdir(FORMS_DIR):
        # Procurar a pasta de ordem de produção
        if "ORDEM DE PRODUCAO" in item.upper() and os.path.isdir(os.path.join(FORMS_DIR, item)):
            ordem_producao_dir = item
            break
    
    if not ordem_producao_dir:
        # Se não encontrar a pasta, mostrar todas as categorias (fallback)
        categories = []
        for item in os.listdir(FORMS_DIR):
            if os.path.isdir(os.path.join(FORMS_DIR, item)):
                categories.append(item)
        categories.sort()
    else:
        # Se encontrar, mostrar apenas essa categoria
        categories = [ordem_producao_dir]
    
    # Adicionar uma mensagem informativa
    flash("Os formulários agora possuem visualização online, impressão direta e download. Clique no botão correspondente para a ação desejada.", "success")
    
    return render_template(
        'forms/index.html',
        title="Formulários de Ordem de Produção",
        categories=categories
    )

@forms_bp.route('/category/<category>')
@login_required
def category(category):
    """Visualizar formulários de uma categoria específica."""
    category_path = os.path.join(FORMS_DIR, category)
    
    # Verificar se a categoria existe
    if not os.path.exists(category_path) or not os.path.isdir(category_path):
        abort(404)
    
    # Listar formulários na categoria
    forms = []
    for root, dirs, files in os.walk(category_path):
        relative_path = os.path.relpath(root, FORMS_DIR)
        
        for file in files:
            if file.startswith('~$') or file.startswith('.'):  # Ignorar arquivos temporários
                continue
                
            file_path = os.path.join(relative_path, file)
            file_ext = os.path.splitext(file)[1].lower()
            
            # Ícone baseado na extensão
            icon = 'fa-file'
            if file_ext in ['.pdf']:
                icon = 'fa-file-pdf'
            elif file_ext in ['.doc', '.docx']:
                icon = 'fa-file-word'
            elif file_ext in ['.xls', '.xlsx']:
                icon = 'fa-file-excel'
            elif file_ext in ['.ppt', '.pptx']:
                icon = 'fa-file-powerpoint'
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
                icon = 'fa-file-image'
            
            forms.append({
                'name': file,
                'path': file_path,
                'icon': icon,
                'date': os.path.getmtime(os.path.join(FORMS_DIR, file_path))
            })
    
    # Ordenar formulários por nome
    forms.sort(key=lambda x: x['name'])
    
    return render_template(
        'forms/category.html',
        title=f"Formulários - {category}",
        category=category,
        forms=forms
    )

@forms_bp.route('/view/<path:file_path>')
@login_required
def view_form(file_path):
    """Visualizar um formulário específico."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    # Obter o tipo MIME do arquivo
    mime_type, _ = mimetypes.guess_type(full_path)
    if not mime_type:
        mime_type = 'application/octet-stream'
    
    # Enviar o arquivo para visualização
    return send_file(
        full_path,
        mimetype=mime_type,
        as_attachment=False,
        download_name=os.path.basename(full_path)
    )

@forms_bp.route('/download/<path:file_path>')
@login_required
def download_form(file_path):
    """Baixar um formulário específico."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    # Enviar o arquivo para download
    return send_file(
        full_path,
        as_attachment=True,
        download_name=os.path.basename(full_path)
    )

@forms_bp.route('/search')
@login_required
def search_forms():
    """Pesquisar formulários."""
    query = request.args.get('q', '').lower()
    if not query or len(query) < 3:
        return jsonify([])
    
    results = []
    
    # Buscar em todas as pastas
    for root, dirs, files in os.walk(FORMS_DIR):
        for file in files:
            if file.startswith('~$') or file.startswith('.'):  # Ignorar arquivos temporários
                continue
                
            if query in file.lower():
                relative_path = os.path.relpath(root, FORMS_DIR)
                file_path = os.path.join(relative_path, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                # Ícone baseado na extensão
                icon = 'fa-file'
                if file_ext in ['.pdf']:
                    icon = 'fa-file-pdf'
                elif file_ext in ['.doc', '.docx']:
                    icon = 'fa-file-word'
                elif file_ext in ['.xls', '.xlsx']:
                    icon = 'fa-file-excel'
                elif file_ext in ['.ppt', '.pptx']:
                    icon = 'fa-file-powerpoint'
                elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
                    icon = 'fa-file-image'
                
                results.append({
                    'name': file,
                    'path': file_path,
                    'category': os.path.basename(root),
                    'icon': icon
                })
    
    # Ordenar resultados por nome
    results.sort(key=lambda x: x['name'])
    
    return jsonify(results)


@forms_bp.route('/api/preset/<int:preset_id>')
@login_required
def get_preset_data(preset_id):
    """API para obter dados de uma predefinição."""
    preset = FormPreset.query.get_or_404(preset_id)
    
    # Verificar se o usuário tem permissão (qualquer usuário pode ver)
    return jsonify({
        'success': True,
        'preset': {
            'id': preset.id,
            'name': preset.name,
            'description': preset.description,
            'form_type': preset.form_type,
            'data': preset.data,
            'is_default': preset.is_default
        }
    })


@forms_bp.route('/api/create-preset/<path:file_path>', methods=['POST'])
@login_required
def create_preset_ajax(file_path):
    """API para criar um preset via AJAX."""
    try:
        full_path = os.path.join(FORMS_DIR, file_path)
        
        # Verificar se o arquivo existe
        if not os.path.exists(full_path) or not os.path.isfile(full_path):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado.'
            }), 404
        
        file_name = os.path.basename(full_path)
        
        # Obter dados da requisição
        data = request.json
        
        if not data or not data.get('name') or not data.get('data'):
            return jsonify({
                'success': False,
                'message': 'Dados incompletos. Nome e campos são obrigatórios.'
            }), 400
        
        preset_name = data.get('name')
        preset_description = data.get('description', '')
        is_default = data.get('is_default', False)
        preset_data = data.get('data', {})
        
        # Verificar se já existe uma predefinição com esse nome
        existing = FormPreset.query.filter_by(
            form_type=file_name,
            name=preset_name
        ).first()
        
        if existing:
            return jsonify({
                'success': False,
                'message': f'Já existe uma predefinição com o nome "{preset_name}".'
            }), 400
        
        # Criar nova predefinição
        preset = FormPreset(
            name=preset_name,
            description=preset_description,
            form_type=file_name,
            file_path=file_path,
            created_by=current_user.id,
            data=preset_data,
            is_default=is_default
        )
        
        # Se esta predefinição é marcada como padrão, desmarcar as demais
        if is_default:
            FormPreset.query.filter_by(
                form_type=file_name,
                is_default=True
            ).update({'is_default': False})
        
        db.session.add(preset)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Predefinição "{preset_name}" criada com sucesso!',
            'preset_id': preset.id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Erro ao criar predefinição: {str(e)}'
        }), 500


@forms_bp.route('/fill/<path:file_path>')
@login_required
def fill_form(file_path):
    """Interface para preencher um formulário."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    file_name = os.path.basename(full_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    # Obter campos disponíveis para preenchimento (baseado no tipo de arquivo)
    fields = get_form_fields(full_path)
    
    # Obter predefinições disponíveis para esse formulário
    presets = FormPreset.query.filter_by(
        form_type=file_name,
        is_active=True
    ).order_by(FormPreset.is_default.desc(), FormPreset.name).all()
    
    return render_template(
        'forms/fill_form.html',
        title=f"Preencher Formulário - {file_name}",
        file_path=file_path,
        file_name=file_name,
        file_ext=file_ext,
        fields=fields,
        presets=presets
    )


@forms_bp.route('/interactive/<path:file_path>')
@login_required
def interactive_form(file_path):
    """Interface interativa para visualizar e preencher formulários."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('forms.index'))
    
    file_name = os.path.basename(full_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    # Obter campos disponíveis para preenchimento
    fields = get_form_fields(full_path)
    
    # Obter predefinições disponíveis para esse formulário
    presets = FormPreset.query.filter_by(
        form_type=file_name,
        is_active=True
    ).order_by(FormPreset.is_default.desc(), FormPreset.name).all()
    
    # Obter campos padronizados disponíveis
    standard_fields = StandardFields.query.filter_by(
        is_active=True
    ).order_by(StandardFields.is_default.desc(), StandardFields.name).all()
    
    # Gerar CSRF token para o template
    csrf_token = generate_csrf()
    
    return render_template(
        'forms/interactive_form.html',
        title=f"Preenchimento Interativo - {file_name}",
        file_path=file_path,
        file_name=file_name,
        file_ext=file_ext,
        fields=fields,
        presets=presets,
        standard_fields=standard_fields,
        csrf_token=csrf_token
    )


@forms_bp.route('/presets/<path:file_path>', methods=['GET'])
@login_required
def list_presets(file_path):
    """Lista todas as predefinições para um formulário específico."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    file_name = os.path.basename(full_path)
    
    # Obter predefinições disponíveis para esse formulário
    presets = FormPreset.query.filter_by(
        form_type=file_name,
        is_active=True
    ).order_by(FormPreset.is_default.desc(), FormPreset.name).all()
    
    return render_template(
        'forms/presets.html',
        title=f"Predefinições - {file_name}",
        file_path=file_path,
        file_name=file_name,
        presets=[preset.to_dict() for preset in presets]
    )


@forms_bp.route('/presets/create/<path:file_path>', methods=['GET', 'POST'])
@login_required
def create_preset(file_path):
    """Criar uma nova predefinição para um formulário."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    file_name = os.path.basename(full_path)
    
    if request.method == 'POST':
        # Obter dados do formulário
        data = request.form.to_dict()
        fields_data = {}
        
        # Verificar campos especiais (não são campos do formulário)
        preset_name = data.pop('preset_name', f'Predefinição para {file_name}')
        preset_description = data.pop('preset_description', '')
        is_default = data.pop('is_default', 'off') == 'on'
        
        # Os demais dados são campos do formulário
        for key, value in data.items():
            if key.startswith('field_'):
                field_id = key.replace('field_', '')
                fields_data[field_id] = value
        
        # Criar uma nova predefinição
        preset = FormPreset(
            name=preset_name,
            description=preset_description,
            form_type=file_name,
            file_path=file_path,
            created_by=current_user.id,
            data=fields_data,
            is_default=is_default
        )
        
        # Se esta predefinição é marcada como padrão, desmarcar as demais
        if is_default:
            FormPreset.query.filter_by(
                form_type=file_name,
                is_default=True
            ).update({'is_default': False})
        
        db.session.add(preset)
        db.session.commit()
        
        flash(f'Predefinição "{preset_name}" criada com sucesso!', 'success')
        return redirect(url_for('forms.list_presets', file_path=file_path))
    
    # Obter campos disponíveis para preenchimento (baseado no tipo de arquivo)
    fields = get_form_fields(full_path)
    
    return render_template(
        'forms/create_preset.html',
        title=f"Nova Predefinição - {file_name}",
        file_path=file_path,
        file_name=file_name,
        fields=fields
    )


@forms_bp.route('/presets/edit/<int:preset_id>', methods=['GET', 'POST'])
@login_required
def edit_preset(preset_id):
    """Editar uma predefinição existente."""
    preset = FormPreset.query.get_or_404(preset_id)
    
    # Verificar se o usuário tem permissão (criador ou administrador)
    if preset.created_by != current_user.id and not current_user.is_admin:
        abort(403)
    
    full_path = os.path.join(FORMS_DIR, preset.file_path)
    
    # Verificar se o arquivo ainda existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('O arquivo original não está mais disponível.', 'warning')
        return redirect(url_for('forms.list_presets', file_path=preset.file_path))
    
    if request.method == 'POST':
        # Obter dados do formulário
        data = request.form.to_dict()
        fields_data = {}
        
        # Verificar campos especiais (não são campos do formulário)
        preset.name = data.pop('preset_name', preset.name)
        preset.description = data.pop('preset_description', preset.description)
        is_default = data.pop('is_default', 'off') == 'on'
        
        # Os demais dados são campos do formulário
        for key, value in data.items():
            if key.startswith('field_'):
                field_id = key.replace('field_', '')
                fields_data[field_id] = value
        
        preset.data = fields_data
        
        # Atualizar status padrão
        if is_default and not preset.is_default:
            # Se esta predefinição agora é marcada como padrão, desmarcar as demais
            FormPreset.query.filter_by(
                form_type=preset.form_type,
                is_default=True
            ).update({'is_default': False})
            preset.is_default = True
        elif not is_default and preset.is_default:
            preset.is_default = False
        
        db.session.commit()
        
        flash(f'Predefinição "{preset.name}" atualizada com sucesso!', 'success')
        return redirect(url_for('forms.list_presets', file_path=preset.file_path))
    
    # Obter campos disponíveis para preenchimento (baseado no tipo de arquivo)
    fields = get_form_fields(full_path)
    
    return render_template(
        'forms/edit_preset.html',
        title=f"Editar Predefinição - {preset.name}",
        preset=preset,
        file_path=preset.file_path,
        file_name=os.path.basename(full_path),
        fields=fields
    )


@forms_bp.route('/presets/delete/<int:preset_id>', methods=['POST'])
@login_required
def delete_preset(preset_id):
    """Excluir uma predefinição."""
    preset = FormPreset.query.get_or_404(preset_id)
    
    # Verificar se o usuário tem permissão (criador ou administrador)
    if preset.created_by != current_user.id and not current_user.is_admin:
        abort(403)
    
    file_path = preset.file_path
    preset_name = preset.name
    
    db.session.delete(preset)
    db.session.commit()
    
    flash(f'Predefinição "{preset_name}" excluída com sucesso!', 'success')
    return redirect(url_for('forms.list_presets', file_path=file_path))


@forms_bp.route('/download_filled/<path:file_path>', methods=['POST'])
@login_required
def download_filled_form(file_path):
    """Baixar um formulário preenchido com os dados fornecidos."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    # Obter os dados para preenchimento
    form_data = {}
    
    # Verificar se está usando uma predefinição ou dados do formulário
    if 'preset_id' in request.form and request.form['preset_id']:
        preset_id = request.form['preset_id']
        preset = FormPreset.query.get_or_404(preset_id)
        form_data = preset.data
    else:
        # Dados enviados diretamente pelo formulário
        for key, value in request.form.items():
            if key.startswith('field_'):
                field_id = key.replace('field_', '')
                form_data[field_id] = value
    
    # Gerar arquivo preenchido
    file_name = os.path.basename(full_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    filled_path = fill_form_with_data(full_path, form_data)
    
    if not filled_path:
        flash('Não foi possível preencher o formulário. Formato não suportado.', 'error')
        return redirect(url_for('forms.fill_form', file_path=file_path))
    
    # Enviar o arquivo preenchido para download
    return send_file(
        filled_path,
        as_attachment=True,
        download_name=f'Preenchido_{file_name}'
    )


@forms_bp.route('/preset/<int:preset_id>/download')
@login_required
def download_preset(preset_id):
    """Baixar um formulário preenchido com uma predefinição específica."""
    preset = FormPreset.query.get_or_404(preset_id)
    full_path = os.path.join(FORMS_DIR, preset.file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('O arquivo original não está mais disponível.', 'warning')
        return redirect(url_for('forms.index'))
    
    # Se usa campos padronizados, incluir esses dados
    form_data = dict(preset.data)  # Copiar para não modificar o original
    
    if preset.use_standard_fields and preset.standard_fields_id:
        standard_fields = StandardFields.query.get(preset.standard_fields_id)
        if standard_fields:
            # Adicionar campos padronizados ao dicionário de dados
            for field in ['empresa', 'produto', 'marca', 'lote', 'responsavel', 'departamento']:
                if hasattr(standard_fields, field) and getattr(standard_fields, field):
                    # Adicionar apenas ao dicionário de preenchimento, não modifica o objeto preset
                    form_data[field] = getattr(standard_fields, field)
    
    filled_path = fill_form_with_data(full_path, form_data)
    
    if not filled_path:
        flash('Não foi possível preencher o formulário. Formato não suportado.', 'error')
        return redirect(url_for('forms.list_presets', file_path=preset.file_path))
    
    file_name = os.path.basename(full_path)
    
    # Enviar o arquivo preenchido para download
    return send_file(
        filled_path,
        as_attachment=True,
        download_name=f'{preset.name}_{file_name}'
    )


# Rotas para gerenciar campos padronizados
@forms_bp.route('/standard-fields')
@login_required
def standard_fields_list():
    """Lista todos os conjuntos de campos padronizados disponíveis."""
    standard_fields = StandardFields.query.filter_by(
        is_active=True
    ).order_by(StandardFields.is_default.desc(), StandardFields.name).all()
    
    return render_template(
        'forms/standard_fields_list.html',
        title="Campos Padronizados",
        standard_fields=standard_fields
    )


@forms_bp.route('/standard-fields/create', methods=['GET', 'POST'])
@login_required
def create_standard_fields():
    """Criar um novo conjunto de campos padronizados."""
    if request.method == 'POST':
        name = request.form.get('name')
        empresa = request.form.get('empresa', 'Zelopack')
        produto = request.form.get('produto')
        marca = request.form.get('marca')
        lote = request.form.get('lote')
        responsavel = request.form.get('responsavel')
        departamento = request.form.get('departamento')
        linha_producao = request.form.get('linha_producao')
        is_default = request.form.get('is_default', 'off') == 'on'
        
        if not name:
            flash('O nome é obrigatório.', 'warning')
            return redirect(url_for('forms.create_standard_fields'))
        
        # Verificar se já existe um conjunto com este nome
        existing = StandardFields.query.filter_by(name=name).first()
        if existing:
            flash(f'Já existe um conjunto de campos com o nome "{name}".', 'warning')
            return redirect(url_for('forms.create_standard_fields'))
        
        # Converter datas se fornecidas
        data_producao = None
        data_validade = None
        
        if request.form.get('data_producao'):
            try:
                data_producao = datetime.datetime.strptime(
                    request.form.get('data_producao'), '%Y-%m-%d').date()
            except ValueError:
                flash('Formato de data de produção inválido.', 'warning')
        
        if request.form.get('data_validade'):
            try:
                data_validade = datetime.datetime.strptime(
                    request.form.get('data_validade'), '%Y-%m-%d').date()
            except ValueError:
                flash('Formato de data de validade inválido.', 'warning')
        
        # Criar novo conjunto de campos padronizados
        std_fields = StandardFields(
            name=name,
            empresa=empresa,
            produto=produto,
            marca=marca,
            lote=lote,
            data_producao=data_producao,
            data_validade=data_validade,
            responsavel=responsavel,
            departamento=departamento,
            linha_producao=linha_producao,
            created_by=current_user.id,
            is_default=is_default
        )
        
        # Se este conjunto está marcado como padrão, desmarcar os demais
        if is_default:
            StandardFields.query.filter_by(is_default=True).update({'is_default': False})
        
        db.session.add(std_fields)
        db.session.commit()
        
        flash(f'Conjunto de campos "{name}" criado com sucesso!', 'success')
        return redirect(url_for('forms.standard_fields_list'))
    
    return render_template(
        'forms/create_standard_fields.html',
        title="Criar Campos Padronizados"
    )


@forms_bp.route('/standard-fields/edit/<int:fields_id>', methods=['GET', 'POST'])
@login_required
def edit_standard_fields(fields_id):
    """Editar um conjunto de campos padronizados existente."""
    std_fields = StandardFields.query.get_or_404(fields_id)
    
    # Verificar se o usuário tem permissão (criador ou administrador)
    if std_fields.created_by != current_user.id and not current_user.is_admin:
        abort(403)
    
    if request.method == 'POST':
        std_fields.name = request.form.get('name', std_fields.name)
        std_fields.empresa = request.form.get('empresa', std_fields.empresa)
        std_fields.produto = request.form.get('produto')
        std_fields.marca = request.form.get('marca')
        std_fields.lote = request.form.get('lote')
        std_fields.responsavel = request.form.get('responsavel')
        std_fields.departamento = request.form.get('departamento')
        std_fields.linha_producao = request.form.get('linha_producao')
        is_default = request.form.get('is_default', 'off') == 'on'
        
        # Converter datas se fornecidas
        if request.form.get('data_producao'):
            try:
                std_fields.data_producao = datetime.datetime.strptime(
                    request.form.get('data_producao'), '%Y-%m-%d').date()
            except ValueError:
                flash('Formato de data de produção inválido.', 'warning')
        else:
            std_fields.data_producao = None
        
        if request.form.get('data_validade'):
            try:
                std_fields.data_validade = datetime.datetime.strptime(
                    request.form.get('data_validade'), '%Y-%m-%d').date()
            except ValueError:
                flash('Formato de data de validade inválido.', 'warning')
        else:
            std_fields.data_validade = None
        
        # Atualizar status padrão
        if is_default and not std_fields.is_default:
            # Se este conjunto agora é marcado como padrão, desmarcar os demais
            StandardFields.query.filter_by(is_default=True).update({'is_default': False})
            std_fields.is_default = True
        elif not is_default and std_fields.is_default:
            std_fields.is_default = False
        
        db.session.commit()
        
        flash(f'Conjunto de campos "{std_fields.name}" atualizado com sucesso!', 'success')
        return redirect(url_for('forms.standard_fields_list'))
    
    return render_template(
        'forms/edit_standard_fields.html',
        title=f"Editar Campos Padronizados - {std_fields.name}",
        fields=std_fields
    )


@forms_bp.route('/standard-fields/delete/<int:fields_id>', methods=['POST'])
@login_required
def delete_standard_fields(fields_id):
    """Excluir um conjunto de campos padronizados."""
    std_fields = StandardFields.query.get_or_404(fields_id)
    
    # Verificar se o usuário tem permissão (criador ou administrador)
    if std_fields.created_by != current_user.id and not current_user.is_admin:
        abort(403)
    
    name = std_fields.name
    
    # Verificar se está sendo usado em predefinições
    presets_using = FormPreset.query.filter_by(standard_fields_id=fields_id).count()
    if presets_using > 0:
        flash(f'Este conjunto de campos está sendo usado em {presets_using} predefinições. Não é possível excluí-lo.', 'warning')
        return redirect(url_for('forms.standard_fields_list'))
    
    db.session.delete(std_fields)
    db.session.commit()
    
    flash(f'Conjunto de campos "{name}" excluído com sucesso!', 'success')
    return redirect(url_for('forms.standard_fields_list'))


@forms_bp.route('/api/standard-fields/<int:fields_id>')
@login_required
def get_standard_fields_data(fields_id):
    """API para obter dados de um conjunto de campos padronizados."""
    std_fields = StandardFields.query.get_or_404(fields_id)
    
    return jsonify({
        'success': True,
        'fields': std_fields.to_dict()
    })


# Funções auxiliares para manipulação de formulários

def get_form_fields(file_path):
    """Extrai os campos disponíveis para preenchimento em um formulário."""
    file_ext = os.path.splitext(file_path)[1].lower()
    fields = []
    
    try:
        if file_ext == '.xlsx' or file_ext == '.xls':
            try:
                # Extrair campos de planilha Excel
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                sheet = workbook.active
                
                if sheet is None:
                    flash('Não foi possível acessar a planilha.', 'warning')
                    return fields
                    
                # Obter o número máximo de linhas e colunas com segurança
                max_row = sheet.max_row or 1
                max_col = sheet.max_column or 1
                
                for row in range(1, max_row + 1):
                    for col in range(1, max_col + 1):
                        try:
                            cell = sheet.cell(row=row, column=col)
                            if cell is None:
                                continue
                            
                            value = cell.value
                            
                            if value and isinstance(value, str) and ('___' in value or '____' in value):
                                # Encontrou um campo para preenchimento (representado por sublinhados)
                                field_id = f"cell_{row}_{col}"
                                sheet_title = getattr(sheet, 'title', 'Planilha')
                                fields.append({
                                    'id': field_id,
                                    'name': f"Campo em {sheet_title} ({get_column_letter(col)}{row})",
                                    'value': ''
                                })
                        except Exception as cell_error:
                            logger.debug(f"Erro ao processar célula ({row},{col}): {cell_error}")
                            continue
            except Exception as excel_error:
                logger.debug(f"Erro ao processar arquivo Excel: {excel_error}")
                fields.append({
                    'id': 'error_field',
                    'name': 'Erro ao ler o arquivo Excel (clique para ver detalhes)',
                    'value': str(excel_error)
                })
            
        elif file_ext == '.docx':
            try:
                # Extrair campos de documento Word
                doc = docx.Document(file_path)
                
                for para_index, para in enumerate(doc.paragraphs):
                    try:
                        text = para.text
                        if '___' in text or '____' in text:
                            # Encontrou um parágrafo com campos para preenchimento
                            field_id = f"para_{para_index}"
                            fields.append({
                                'id': field_id,
                                'name': f"Campo no parágrafo {para_index + 1}: {text[:30]}...",
                                'value': ''
                            })
                    except Exception as para_error:
                        logger.debug(f"Erro ao processar parágrafo {para_index}: {para_error}")
                        continue
            except Exception as docx_error:
                logger.debug(f"Erro ao processar arquivo Word: {docx_error}")
                fields.append({
                    'id': 'error_field',
                    'name': 'Erro ao ler o arquivo Word (clique para ver detalhes)',
                    'value': str(docx_error)
                })
            
        elif file_ext == '.pdf':
            try:
                # Extrair campos de PDF (mais complexo)
                # Implementação básica para detecção de campos
                reader = PyPDF2.PdfReader(file_path)
                form_fields = reader.get_fields()
                
                if form_fields:
                    # PDF tem campos de formulário
                    for field_name, field_value in form_fields.items():
                        fields.append({
                            'id': field_name,
                            'name': field_name,
                            'value': ''
                        })
                else:
                    # PDF não tem campos de formulário, tentar identificar por texto
                    for page_num in range(len(reader.pages)):
                        try:
                            page = reader.pages[page_num]
                            text = page.extract_text()
                            
                            if text and ('___' in text or '____' in text):
                                # Encontrou texto com campos potenciais
                                field_id = f"pdf_page_{page_num}"
                                fields.append({
                                    'id': field_id,
                                    'name': f"Campo na página {page_num + 1}",
                                    'value': ''
                                })
                        except Exception as page_error:
                            logger.debug(f"Erro ao processar página {page_num} do PDF: {page_error}")
                            continue
            except Exception as pdf_error:
                logger.debug(f"Erro ao processar arquivo PDF: {pdf_error}")
                fields.append({
                    'id': 'error_field',
                    'name': 'Erro ao ler o arquivo PDF (clique para ver detalhes)',
                    'value': str(pdf_error)
                })
        else:
            # Formato não suportado
            fields.append({
                'id': 'unsupported_format',
                'name': f'Formato não suportado: {file_ext}',
                'value': 'Este tipo de arquivo não pode ser processado automaticamente.'
            })
    except Exception as e:
        logger.debug(f"Erro ao extrair campos do formulário: {e}")
        fields.append({
            'id': 'general_error',
            'name': 'Erro ao processar o formulário (clique para ver detalhes)',
            'value': str(e)
        })
    
    # Se não encontrou nenhum campo, adicionar uma mensagem informativa
    if not fields:
        fields.append({
            'id': 'no_fields',
            'name': 'Nenhum campo detectado',
            'value': 'O sistema não conseguiu detectar campos para preenchimento neste documento.'
        })
    
    return fields


def fill_form_with_data(file_path, form_data):
    """Preenche um formulário com os dados fornecidos e retorna o caminho do arquivo preenchido."""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Criar diretório temporário para o formulário preenchido
    temp_dir = tempfile.mkdtemp()
    file_name = os.path.basename(file_path)
    output_path = os.path.join(temp_dir, f'filled_{file_name}')
    
    try:
        # Identificar e ignorar campos que não são para preenchimento de dados
        valid_form_data = {}
        for field_id, value in form_data.items():
            if field_id in ['error_field', 'general_error', 'no_fields', 'unsupported_format']:
                continue
            valid_form_data[field_id] = value
        
        if file_ext == '.xlsx' or file_ext == '.xls':
            try:
                # Preencher planilha Excel
                workbook = openpyxl.load_workbook(file_path)
                sheet = workbook.active
                
                if sheet is None:
                    raise ValueError('Não foi possível acessar a planilha.')
                
                # Aplicar os dados aos campos identificados
                for field_id, value in valid_form_data.items():
                    if field_id.startswith('cell_'):
                        parts = field_id.split('_')
                        if len(parts) >= 3:
                            try:
                                _, row, col = parts
                                row, col = int(row), int(col)
                                
                                cell = sheet.cell(row=row, column=col)
                                if cell is None:
                                    continue
                                
                                # Verificar se é célula vazia ou tem marcadores de campo
                                original_value = cell.value
                                if original_value is None or (isinstance(original_value, str) and 
                                                             ('___' in original_value or '____' in original_value)):
                                    cell.value = value
                                    
                                    # Aplicar estilo para destacar o campo preenchido
                                    cell.font = Font(bold=True, color="0000FF")
                            except Exception as cell_error:
                                logger.debug(f"Erro ao preencher célula ({field_id}): {cell_error}")
                                continue
                
                # Salvar a planilha preenchida
                workbook.save(output_path)
            except Exception as excel_error:
                logger.debug(f"Erro ao processar arquivo Excel para preenchimento: {excel_error}")
                raise ValueError(f"Não foi possível preencher o formulário Excel: {str(excel_error)}")
            
        elif file_ext == '.docx':
            try:
                # Preencher documento Word
                doc = docx.Document(file_path)
                
                # Aplicar os dados aos campos identificados
                for field_id, value in valid_form_data.items():
                    if field_id.startswith('para_'):
                        try:
                            parts = field_id.split('_')
                            if len(parts) >= 2:
                                para_index = int(parts[1])
                                if para_index < len(doc.paragraphs):
                                    para = doc.paragraphs[para_index]
                                    text = para.text
                                    
                                    # Substituir campos em branco pelo valor
                                    new_text = text.replace('_____', value).replace('____', value).replace('___', value)
                                    
                                    # Limpar o parágrafo e adicionar o texto substituído
                                    para.clear()
                                    para.add_run(new_text)
                        except Exception as para_error:
                            logger.debug(f"Erro ao preencher parágrafo {field_id}: {para_error}")
                            continue
                
                # Salvar o documento preenchido
                doc.save(output_path)
            except Exception as docx_error:
                logger.debug(f"Erro ao processar arquivo Word para preenchimento: {docx_error}")
                raise ValueError(f"Não foi possível preencher o documento Word: {str(docx_error)}")
            
        elif file_ext == '.pdf':
            try:
                # Para PDFs, criamos um novo PDF com o texto sobreposto
                # (abordagem simples, funcionalidade limitada)
                reader = PyPDF2.PdfReader(file_path)
                writer = PyPDF2.PdfWriter()
                
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    writer.add_page(page)
                
                # Se o PDF tem campos de formulário
                form_fields = reader.get_fields()
                if form_fields:
                    # Preencher campos de formulário
                    update_fields = {}
                    for field_id, value in valid_form_data.items():
                        if field_id in form_fields:
                            update_fields[field_id] = value
                    
                    if update_fields:
                        for page_num in range(len(reader.pages)):
                            try:
                                writer.update_page_form_field_values(page_num, update_fields)
                            except Exception as page_error:
                                logger.debug(f"Erro ao preencher página {page_num} do PDF: {page_error}")
                                continue
                
                # Salvar o PDF preenchido
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
            except Exception as pdf_error:
                logger.debug(f"Erro ao processar arquivo PDF para preenchimento: {pdf_error}")
                raise ValueError(f"Não foi possível preencher o documento PDF: {str(pdf_error)}")
                
        else:
            # Formato não suportado
            raise ValueError(f"Formato de arquivo não suportado: {file_ext}")
    
    except Exception as e:
        logger.debug(f"Erro ao preencher formulário: {e}")
        try:
            # Limpar recursos temporários
            shutil.rmtree(temp_dir)
        except:
            pass
        return None
    
    return output_path


# A função get_column_letter já está importada de openpyxl.utils