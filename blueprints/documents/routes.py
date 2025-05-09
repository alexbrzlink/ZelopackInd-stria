from flask import render_template, redirect, url_for, request, flash, jsonify, current_app, send_file, abort
from flask_login import login_required, current_user
from flask_wtf.csrf import generate_csrf
from sqlalchemy import desc, func, or_
from werkzeug.utils import secure_filename
import os
import datetime
import io
import tempfile
import logging

# Importações para manipulação de imagens e documentos
from PIL import Image, ImageDraw, ImageFont

# Importações opcionais (usadas na função de preview)
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None
    letter = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

from app import db
from models import TechnicalDocument, DocumentAttachment, User
from blueprints.documents import documents_bp
from blueprints.documents.forms import DocumentForm, DocumentSearchForm

# Configuração para criar miniaturas de imagens
THUMBNAIL_SIZE = (200, 200)

@documents_bp.route('/')
@login_required
def index():
    """Página principal do módulo de documentos técnicos."""
    search_form = DocumentSearchForm()
    query = TechnicalDocument.query.order_by(TechnicalDocument.upload_date.desc())
    
    # Filtrar por tipo de documento ativo por padrão
    query = query.filter(TechnicalDocument.status == 'ativo')
    
    # Obter contagem por tipo de documento
    doc_stats = db.session.query(
        TechnicalDocument.document_type, 
        func.count(TechnicalDocument.id)
    ).group_by(TechnicalDocument.document_type).all()
    
    # Formatar estatísticas
    stats = {}
    for doc_type, count in doc_stats:
        if doc_type == 'pop':
            stats['POPs'] = count
        elif doc_type == 'ficha_tecnica':
            stats['Fichas Técnicas'] = count
        elif doc_type == 'certificado':
            stats['Certificados'] = count
        elif doc_type == 'instrucao':
            stats['Instruções'] = count
        elif doc_type == 'planilha':
            stats['Planilhas'] = count
        elif doc_type == 'manual':
            stats['Manuais'] = count
        elif doc_type == 'formulario':
            stats['Formulários'] = count
        else:
            stats['Outros'] = stats.get('Outros', 0) + count
    
    # Obter documentos recentes
    recent_docs = query.limit(10).all()
    
    # Obter documentos por categoria para formulários
    formularios = {}
    
    # Importar as bibliotecas necessárias
    import os
    
    # Definir o diretório de formulários
    FORMS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'extracted_forms')
    
    # Mapeamento de categorias para pasta real
    mapeamento_categorias = {
        'blender': 'FORMULÁRIOS BLENDER',
        'laboratorio': 'FORMULÁRIOS LABORATÓRIO',
        'portaria': 'FORMULÁRIOS PORTARIA', 
        'qualidade': 'FORMULÁRIOS QUALIDADE',
        'tba': 'FORMULÁRIOS TBA',
    }
    
    # Categorias que queremos mostrar
    categorias = ['blender', 'laboratorio', 'portaria', 'qualidade', 'tba']
    
    for categoria in categorias:
        # Nome da categoria em maiúsculas
        nome_categoria = categoria.upper()
        
        # Pasta real no sistema de arquivos
        pasta_real = mapeamento_categorias.get(categoria)
        pasta_completa = os.path.join(FORMS_DIR, pasta_real) if pasta_real else None
        
        # Lista para armazenar os documentos virtuais
        docs_virtuais = []
        
        # Verificar se a pasta existe
        if pasta_real and pasta_completa and os.path.exists(pasta_completa):
            # Obter arquivos na pasta
            for root, dirs, files in os.walk(pasta_completa):
                for arquivo in files:
                    # Ignorar temporários e ocultos
                    if arquivo.startswith('~$') or arquivo.startswith('.'):
                        continue
                        
                    # Caminho relativo ao diretório base
                    caminho_relativo = os.path.relpath(os.path.join(root, arquivo), FORMS_DIR)
                    caminho_completo = os.path.join(FORMS_DIR, caminho_relativo)
                    
                    # Extensão do arquivo
                    extensao = os.path.splitext(arquivo)[1].lower()
                    
                    # Data de modificação
                    data_mod = os.path.getmtime(caminho_completo)
                    
                    # Determinar ícone baseado na extensão
                    if extensao in ['.pdf']:
                        icone = 'fa-file-pdf'
                    elif extensao in ['.doc', '.docx']:
                        icone = 'fa-file-word'
                    elif extensao in ['.xls', '.xlsx']:
                        icone = 'fa-file-excel'
                    elif extensao in ['.ppt', '.pptx']:
                        icone = 'fa-file-powerpoint'
                    elif extensao in ['.jpg', '.jpeg', '.png', '.gif']:
                        icone = 'fa-file-image'
                    else:
                        icone = 'fa-file'
                    
                    # Criar documento virtual
                    doc_virtual = {
                        'id': f"file_{pasta_real}_{arquivo}".replace(" ", "_"),
                        'title': arquivo,
                        'file_path': caminho_relativo,
                        'description': f"Formulário {pasta_real}",
                        'created_at': data_mod,
                        'icon_class': icone,
                        'is_virtual': True
                    }
                    
                    docs_virtuais.append(doc_virtual)
            
            # Adicionar documentos do banco de dados
            docs_bd = TechnicalDocument.query.filter_by(
                document_type='formulario',
                category=categoria,
                status='ativo'
            ).order_by(TechnicalDocument.title).all()
            
            # Se tiver documentos, adicionar à lista
            if docs_virtuais or docs_bd:
                formularios[nome_categoria] = {
                    'documentos_bd': docs_bd,
                    'documentos_virtuais': docs_virtuais
                }
    
    return render_template(
        'documents/index.html',
        title='Documentos Técnicos',
        recent_docs=recent_docs,
        search_form=search_form,
        stats=stats,
        formularios=formularios
    )


@documents_bp.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_document():
    """Upload de um novo documento técnico."""
    form = DocumentForm()
    
    if form.validate_on_submit():
        # Salvar arquivo principal
        document_file = form.document_file.data
        upload_folder = current_app.config['UPLOAD_FOLDER']
        document_folder = os.path.join(upload_folder, 'documents')
        os.makedirs(document_folder, exist_ok=True)
        
        # Gerar nome de arquivo seguro com timestamp
        secure_name = secure_filename(document_file.filename)
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{secure_name}"
        file_path = os.path.join(document_folder, filename)
        
        # Salvar arquivo no sistema
        document_file.save(file_path)
        file_size = os.path.getsize(file_path)
        file_type = os.path.splitext(filename)[1][1:].lower()
        
        # Criar registro no banco de dados
        document = TechnicalDocument(
            title=form.title.data,
            description=form.description.data,
            document_type=form.document_type.data,
            category=form.category.data if form.category.data else None,
            revision=form.revision.data,
            valid_until=form.valid_until.data,
            author=form.author.data,
            tags=form.tags.data,
            status=form.status.data,
            restricted_access=form.restricted_access.data,
            filename=filename,
            original_filename=document_file.filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            uploaded_by=current_user.id,
            version=1
        )
        
        db.session.add(document)
        db.session.commit()
        
        # Processar anexos se houver
        if form.attachments.data:
            attachments_folder = os.path.join(document_folder, f"doc_{document.id}_attachments")
            os.makedirs(attachments_folder, exist_ok=True)
            
            for attachment in form.attachments.data:
                if attachment.filename:
                    secure_name = secure_filename(attachment.filename)
                    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                    attach_filename = f"{timestamp}_{secure_name}"
                    attach_path = os.path.join(attachments_folder, attach_filename)
                    
                    attachment.save(attach_path)
                    attach_size = os.path.getsize(attach_path)
                    attach_type = os.path.splitext(attach_filename)[1][1:].lower()
                    
                    # Criar registro para o anexo
                    doc_attachment = DocumentAttachment(
                        document_id=document.id,
                        filename=attach_filename,
                        original_filename=attachment.filename,
                        file_path=attach_path,
                        file_type=attach_type,
                        file_size=attach_size,
                        description=f"Anexo de {attachment.filename}",
                        uploaded_by=current_user.id
                    )
                    db.session.add(doc_attachment)
            
            db.session.commit()
        
        flash('Documento técnico enviado com sucesso!', 'success')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    return render_template(
        'documents/upload.html',
        title='Upload de Documento Técnico',
        form=form
    )


@documents_bp.route('/view/<int:document_id>')
@login_required
def view_document(document_id):
    """Visualizar um documento técnico."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão para documentos restritos
    if document.restricted_access and current_user.role != 'admin':
        flash('Você não tem permissão para acessar este documento.', 'warning')
        return redirect(url_for('documents.index'))
    
    # Verificar se é para visualização online
    online_view = request.args.get('online', False)
    
    # Obter anexos
    attachments = DocumentAttachment.query.filter_by(document_id=document.id).all()
    
    # Verificar se existem outras versões
    versions = []
    if document.parent_id:
        # Se este documento é uma versão, buscar o original e outras versões
        original = TechnicalDocument.query.get(document.parent_id)
        if original:
            versions.append(original)
        other_versions = TechnicalDocument.query.filter_by(parent_id=document.parent_id).all()
        for v in other_versions:
            if v.id != document.id:
                versions.append(v)
    else:
        # Se este é o original, buscar suas versões
        versions = TechnicalDocument.query.filter_by(parent_id=document.id).all()
    
    # Se for visualização online e o arquivo existir
    if online_view and document.file_path and os.path.exists(document.file_path):
        file_ext = os.path.splitext(document.file_path)[1].lower()
        
        # Para arquivos HTML, exibir diretamente no iframe
        if file_ext == '.html':
            with open(document.file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            return render_template(
                'documents/view_online.html',
                title=f'Visualização Online: {document.title}',
                document=document,
                html_content=html_content
            )
        
        # Para PDFs, exibir no iframe usando o viewer nativo do navegador
        elif file_ext == '.pdf':
            return render_template(
                'documents/view_online.html',
                title=f'Visualização Online: {document.title}',
                document=document,
                pdf_path=url_for('documents.download_document', document_id=document.id)
            )
        
        # Para outros tipos, tentar embutir ou oferecer download
        else:
            flash('Visualização online não disponível para este tipo de arquivo. Faça o download para visualizar.', 'info')
    
    # Renderizar a visualização padrão
    return render_template(
        'documents/view.html',
        title=f'Documento: {document.title}',
        document=document,
        attachments=attachments,
        versions=versions
    )


@documents_bp.route('/editor', methods=['GET', 'POST'])
@login_required
def editor():
    """Editor de documentos técnicos em tempo real."""
    # Obter empresas, produtos e marcas do banco de dados
    # Para fins de demonstração, usaremos alguns exemplos
    empresas = ['Zelopack', 'Fornecedor A', 'Fornecedor B', 'Cliente X', 'Cliente Y']
    produtos = ['Suco de Laranja', 'Suco de Uva', 'Néctar de Maracujá', 'Refresco de Limão', 'Água de Coco']
    marcas = ['Marca A', 'Marca B', 'Marca C', 'Marca Premium', 'Marca Econômica']
    
    # Verificar se há dados de fornecedores disponíveis
    fornecedores_df = None
    try:
        import pandas as pd
        import os
        
        fornecedores_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'fornecedores.xlsx')
        if os.path.exists(fornecedores_path):
            fornecedores_df = pd.read_excel(fornecedores_path)
    except Exception as e:
        logging.error(f"Erro ao carregar dados de fornecedores: {str(e)}")
    
    # Processar envio de documentos
    if request.method == 'POST':
        content = request.form.get('content', '')
        format_type = request.form.get('format', 'pdf')
        
        # Criar documento temporário
        with tempfile.NamedTemporaryFile(suffix=f'.{format_type}', delete=False) as temp_file:
            temp_path = temp_file.name
            
            if format_type == 'pdf':
                # Criar PDF usando ReportLab se disponível
                if canvas and letter:
                    pdf = canvas.Canvas(temp_path, pagesize=letter)
                    lines = content.split('\n')
                    y = 750  # Posição inicial Y
                    for line in lines:
                        pdf.drawString(50, y, line)
                        y -= 15  # Espaçamento entre linhas
                    pdf.save()
                else:
                    # Fallback simples se ReportLab não estiver disponível
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            elif format_type == 'word':
                # Criar arquivo DOCX simples
                try:
                    from docx import Document
                    doc = Document()
                    for paragraph in content.split('\n\n'):
                        doc.add_paragraph(paragraph)
                    doc.save(temp_path)
                except ImportError:
                    # Fallback: salvar como texto simples
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            elif format_type == 'excel':
                # Criar arquivo Excel simples
                try:
                    import pandas as pd
                    # Converter texto para linhas e colunas CSV
                    data = []
                    for line in content.split('\n'):
                        if line.strip():
                            data.append(line.split(','))
                    df = pd.DataFrame(data)
                    df.to_excel(temp_path, index=False, header=False)
                except ImportError:
                    # Fallback: salvar como CSV
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            else:
                # Formato de texto simples
                with open(temp_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            # Retornar o arquivo para download
            return send_file(
                temp_path,
                as_attachment=True,
                download_name=f"documento_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.{format_type}"
            )
    
    return render_template(
        'documents/editor.html',
        empresas=empresas,
        produtos=produtos,
        marcas=marcas,
        fornecedores_df=fornecedores_df
    )


@documents_bp.route('/edit/<int:document_id>', methods=['GET', 'POST'])
@login_required
def edit_document(document_id):
    """Editar um documento técnico."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão
    if document.uploaded_by != current_user.id and current_user.role != 'admin':
        flash('Você não tem permissão para editar este documento.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    form = DocumentForm(obj=document)
    
    # Inicialmente não exigimos novo upload ao editar
    form.document_file.validators = []
    
    if form.validate_on_submit():
        # Atualizar campos básicos
        document.title = form.title.data
        document.description = form.description.data
        document.document_type = form.document_type.data
        document.category = form.category.data if form.category.data else None
        document.revision = form.revision.data
        document.valid_until = form.valid_until.data
        document.author = form.author.data
        document.tags = form.tags.data
        document.status = form.status.data
        document.restricted_access = form.restricted_access.data
        
        # Processar novo arquivo principal (se fornecido)
        if form.document_file.data and form.document_file.data.filename:
            document_file = form.document_file.data
            upload_folder = current_app.config['UPLOAD_FOLDER']
            document_folder = os.path.join(upload_folder, 'documents')
            
            # Gerar nome de arquivo seguro com timestamp
            secure_name = secure_filename(document_file.filename)
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{timestamp}_{secure_name}"
            file_path = os.path.join(document_folder, filename)
            
            # Salvar arquivo no sistema
            document_file.save(file_path)
            
            # Atualizar informações do arquivo
            document.filename = filename
            document.original_filename = document_file.filename
            document.file_path = file_path
            document.file_type = os.path.splitext(filename)[1][1:].lower()
            document.file_size = os.path.getsize(file_path)
        
        # Processar novos anexos (se fornecidos)
        if form.attachments.data and any(attachment.filename for attachment in form.attachments.data):
            upload_folder = current_app.config['UPLOAD_FOLDER']
            attachments_folder = os.path.join(upload_folder, 'documents', f"doc_{document.id}_attachments")
            os.makedirs(attachments_folder, exist_ok=True)
            
            for attachment in form.attachments.data:
                if attachment.filename:
                    secure_name = secure_filename(attachment.filename)
                    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                    attach_filename = f"{timestamp}_{secure_name}"
                    attach_path = os.path.join(attachments_folder, attach_filename)
                    
                    attachment.save(attach_path)
                    attach_size = os.path.getsize(attach_path)
                    attach_type = os.path.splitext(attach_filename)[1][1:].lower()
                    
                    # Criar registro para o anexo
                    doc_attachment = DocumentAttachment(
                        document_id=document.id,
                        filename=attach_filename,
                        original_filename=attachment.filename,
                        file_path=attach_path,
                        file_type=attach_type,
                        file_size=attach_size,
                        description=f"Anexo de {attachment.filename}",
                        uploaded_by=current_user.id
                    )
                    db.session.add(doc_attachment)
        
        db.session.commit()
        flash('Documento atualizado com sucesso!', 'success')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    return render_template(
        'documents/edit.html',
        title=f'Editar: {document.title}',
        form=form,
        document=document
    )


@documents_bp.route('/new-version/<int:document_id>', methods=['GET', 'POST'])
@login_required
def new_version(document_id):
    """Criar uma nova versão de um documento existente."""
    original_doc = TechnicalDocument.query.get_or_404(document_id)
    
    # Inicializar formulário com dados do documento original
    form = DocumentForm(obj=original_doc)
    
    if form.validate_on_submit():
        # Processar o arquivo principal da nova versão
        document_file = form.document_file.data
        upload_folder = current_app.config['UPLOAD_FOLDER']
        document_folder = os.path.join(upload_folder, 'documents')
        
        # Gerar nome de arquivo seguro com timestamp
        secure_name = secure_filename(document_file.filename)
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{secure_name}"
        file_path = os.path.join(document_folder, filename)
        
        # Salvar arquivo no sistema
        document_file.save(file_path)
        file_size = os.path.getsize(file_path)
        file_type = os.path.splitext(filename)[1][1:].lower()
        
        # Determinar qual será o documento pai (original)
        parent_id = original_doc.parent_id or original_doc.id
        
        # Determinar a nova versão
        latest_version = db.session.query(func.max(TechnicalDocument.version)).filter(
            or_(TechnicalDocument.id == parent_id, TechnicalDocument.parent_id == parent_id)
        ).scalar() or 1
        
        # Criar nova versão no banco de dados
        new_doc = TechnicalDocument(
            title=form.title.data,
            description=form.description.data,
            document_type=form.document_type.data,
            category=form.category.data if form.category.data else None,
            revision=form.revision.data,
            valid_until=form.valid_until.data,
            author=form.author.data,
            tags=form.tags.data,
            status=form.status.data,
            restricted_access=form.restricted_access.data,
            filename=filename,
            original_filename=document_file.filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            uploaded_by=current_user.id,
            parent_id=parent_id,
            version=latest_version + 1
        )
        
        db.session.add(new_doc)
        db.session.commit()
        
        # Processar anexos se houver
        if form.attachments.data:
            attachments_folder = os.path.join(document_folder, f"doc_{new_doc.id}_attachments")
            os.makedirs(attachments_folder, exist_ok=True)
            
            for attachment in form.attachments.data:
                if attachment.filename:
                    secure_name = secure_filename(attachment.filename)
                    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                    attach_filename = f"{timestamp}_{secure_name}"
                    attach_path = os.path.join(attachments_folder, attach_filename)
                    
                    attachment.save(attach_path)
                    attach_size = os.path.getsize(attach_path)
                    attach_type = os.path.splitext(attach_filename)[1][1:].lower()
                    
                    # Criar registro para o anexo
                    doc_attachment = DocumentAttachment(
                        document_id=new_doc.id,
                        filename=attach_filename,
                        original_filename=attachment.filename,
                        file_path=attach_path,
                        file_type=attach_type,
                        file_size=attach_size,
                        description=f"Anexo de {attachment.filename}",
                        uploaded_by=current_user.id
                    )
                    db.session.add(doc_attachment)
            
            db.session.commit()
        
        flash('Nova versão do documento criada com sucesso!', 'success')
        return redirect(url_for('documents.view_document', document_id=new_doc.id))
    
    # Inicializar informações de versão
    if original_doc.parent_id:
        parent_doc = TechnicalDocument.query.get(original_doc.parent_id)
        current_version = original_doc.version
    else:
        parent_doc = original_doc
        current_version = original_doc.version or 1
    
    return render_template(
        'documents/new_version.html',
        title=f'Nova Versão: {original_doc.title}',
        form=form,
        original_doc=original_doc,
        parent_doc=parent_doc,
        current_version=current_version
    )


@documents_bp.route('/print/<int:document_id>')
@login_required
def print_document(document_id):
    """Versão para impressão de um documento técnico."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão para documentos restritos
    if document.restricted_access and current_user.role != 'admin':
        flash('Você não tem permissão para acessar este documento.', 'warning')
        return redirect(url_for('documents.index'))
    
    if os.path.exists(document.file_path):
        file_ext = os.path.splitext(document.file_path)[1].lower()
        
        # Para PDFs, abrir diretamente para impressão
        if file_ext == '.pdf':
            return send_file(
                document.file_path,
                as_attachment=False,
                download_name=document.original_filename
            )
        # Para arquivos HTML, exibir diretamente no iframe
        elif file_ext == '.html':
            with open(document.file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            return render_template(
                'documents/print_view.html',
                title=f'Impressão: {document.title}',
                document=document,
                html_content=html_content
            )
        else:
            # Para outros tipos de arquivo, informar que não é possível imprimir diretamente
            flash('Este tipo de arquivo não pode ser impresso diretamente pelo navegador. Faça o download para imprimir.', 'info')
            return redirect(url_for('documents.view_document', document_id=document.id))
    else:
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.view_document', document_id=document.id))


@documents_bp.route('/view-file/<path:file_path>')
@login_required
def view_virtual_document(file_path):
    """Visualizar um documento virtual (arquivo do sistema)."""
    # Montar o caminho completo
    FORMS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'extracted_forms')
    
    # Garantir que o caminho é seguro (dentro do diretório permitido)
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.index'))
    
    # Verificar a extensão do arquivo
    file_ext = os.path.splitext(full_path)[1].lower()
    file_name = os.path.basename(full_path)
    
    # Determinar o modo de visualização
    online_view = request.args.get('online', '')
    
    # Verificar se a visualização online foi solicitada
    if online_view and online_view.lower() in ('true', '1', 'yes'):
        # Cria um objeto de documento virtual com informações do arquivo
        creation_time = datetime.datetime.fromtimestamp(os.path.getctime(full_path))
        modification_time = datetime.datetime.fromtimestamp(os.path.getmtime(full_path))
        
        doc_info = {
            'title': file_name,
            'file_type': file_ext[1:] if file_ext.startswith('.') else file_ext,
            'updated_at': modification_time,
            'created_at': creation_time,
            'author': 'Sistema Zelopack',
            'category': os.path.dirname(file_path).split('/')[-1] if '/' in file_path else '',
            'revision': '1.0'
        }
        
        # Link para download do arquivo
        download_link = url_for('documents.download_virtual_document', file_path=file_path)
        
        # Para PDFs, exibir no iframe usando o viewer nativo do navegador
        if file_ext == '.pdf':
            file_url = url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false')
            return render_template(
                'documents/view_online.html',
                title=f'Visualização Online: {file_name}',
                document=doc_info,
                pdf_path=file_url,
                download_link=download_link,
                virtual_document=True
            )
        # Para arquivos HTML, exibir diretamente no iframe
        elif file_ext == '.html':
            with open(full_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            return render_template(
                'documents/view_online.html',
                title=f'Visualização Online: {file_name}',
                document=doc_info,
                html_content=html_content,
                download_link=download_link,
                virtual_document=True
            )
        # Para arquivos de imagem
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
            file_url = url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false')
            return render_template(
                'documents/view_online.html',
                title=f'Visualização Online: {file_name}',
                document=doc_info,
                image_path=file_url,
                download_link=download_link,
                virtual_document=True
            )
        # Para planilhas Excel
        elif file_ext in ['.xlsx', '.xls']:
            # Para arquivos Excel, usamos o serviço de visualização do Microsoft Office Online
            file_url = url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false', _external=True)
            # Codificar a URL para o serviço de visualização do Office Online
            # Obs: isso não funciona em localhost, mas funciona quando hospedado
            office_viewer_url = f"https://view.officeapps.live.com/op/embed.aspx?src={file_url}"
            return render_template(
                'documents/view_online.html',
                title=f'Visualização Online: {file_name}',
                document=doc_info,
                office_url=office_viewer_url,
                download_link=download_link,
                virtual_document=True
            )
        else:
            # Para outros tipos, tentar embutir ou oferecer download
            flash('Visualização online não disponível para este tipo de arquivo. Faça o download para visualizar.', 'info')
            return redirect(url_for('documents.download_virtual_document', file_path=file_path))
    else:
        # Download direto
        return redirect(url_for('documents.download_virtual_document', file_path=file_path))


@documents_bp.route('/download-file/<path:file_path>')
@login_required
def download_virtual_document(file_path):
    """Download de um documento virtual (arquivo do sistema)."""
    # Montar o caminho completo
    FORMS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'extracted_forms')
    
    # Garantir que o caminho é seguro (dentro do diretório permitido)
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.index'))
    
    # Determinar se deve baixar como anexo ou visualizar no navegador
    as_attachment_param = request.args.get('as_attachment', 'true')
    as_attachment = as_attachment_param.lower() in ('true', '1', 'yes')
    
    try:
        return send_file(
            full_path,
            as_attachment=as_attachment,
            download_name=os.path.basename(full_path)
        )
    except Exception as e:
        flash(f'Erro ao baixar o arquivo: {str(e)}', 'danger')
        return redirect(url_for('documents.index'))


@documents_bp.route('/print-file/<path:file_path>')
@login_required
def print_virtual_document(file_path):
    """Versão para impressão de um documento virtual (arquivo do sistema)."""
    # Montar o caminho completo
    FORMS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'extracted_forms')
    
    # Garantir que o caminho é seguro (dentro do diretório permitido)
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.index'))
    
    file_ext = os.path.splitext(full_path)[1].lower()
    file_name = os.path.basename(full_path)
    creation_time = datetime.datetime.fromtimestamp(os.path.getctime(full_path))
    modification_time = datetime.datetime.fromtimestamp(os.path.getmtime(full_path))
    
    # Informações do documento para o template
    doc_info = {
        'title': file_name,
        'revision': '1.0',
        'created_at': creation_time,
        'modified_at': modification_time,
        'author': 'Sistema Zelopack',
        'category': os.path.dirname(file_path).split('/')[-1] if '/' in file_path else ''
    }
    
    # Para PDFs, criar uma visualização de impressão com iframe
    if file_ext == '.pdf':
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            pdf_path=url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false'),
            timestamp_to_date=lambda: datetime.datetime.now().strftime('%d/%m/%Y %H:%M')
        )
    # Para arquivos HTML, exibir diretamente no iframe
    elif file_ext == '.html':
        with open(full_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            html_content=html_content,
            timestamp_to_date=lambda: datetime.datetime.now().strftime('%d/%m/%Y %H:%M')
        )
    # Para arquivos Excel
    elif file_ext in ['.xlsx', '.xls']:
        try:
            import pandas as pd
            # Tentar ler a planilha Excel
            df = pd.read_excel(full_path)
            # Converter para HTML para exibir na página de impressão
            table_html = df.to_html(classes='table table-bordered table-striped', index=False)
            return render_template(
                'documents/print_view.html',
                title=f'Impressão: {file_name}',
                document=doc_info,
                html_content=table_html,
                timestamp_to_date=lambda: datetime.datetime.now().strftime('%d/%m/%Y %H:%M')
            )
        except Exception as e:
            # Se não conseguir ler a planilha, oferecer link para visualização online
            return render_template(
                'documents/print_view.html',
                title=f'Impressão: {file_name}',
                document=doc_info,
                error_message=f"Não foi possível gerar uma visualização para impressão deste arquivo. Erro: {str(e)}",
                download_link=url_for('documents.download_virtual_document', file_path=file_path),
                timestamp_to_date=lambda: datetime.datetime.now().strftime('%d/%m/%Y %H:%M')
            )
    # Para imagens, exibir a imagem na página de impressão
    elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            image_path=url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false'),
            timestamp_to_date=lambda: datetime.datetime.now().strftime('%d/%m/%Y %H:%M')
        )
    else:
        # Para outros tipos de arquivo, informar que não é possível imprimir diretamente
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            error_message="Este tipo de arquivo não pode ser convertido automaticamente para impressão.",
            download_link=url_for('documents.download_virtual_document', file_path=file_path),
            timestamp_to_date=lambda: datetime.datetime.now().strftime('%d/%m/%Y %H:%M')
        )


@documents_bp.route('/download/<int:document_id>')
@login_required
def download_document(document_id):
    """Download de um documento técnico."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão para documentos restritos
    if document.restricted_access and current_user.role != 'admin':
        flash('Você não tem permissão para baixar este documento.', 'warning')
        return redirect(url_for('documents.index'))
    
    if os.path.exists(document.file_path):
        return send_file(
            document.file_path,
            as_attachment=True,
            download_name=document.original_filename
        )
    else:
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.view_document', document_id=document.id))


@documents_bp.route('/download-attachment/<int:attachment_id>')
@login_required
def download_attachment(attachment_id):
    """Download de um anexo de documento."""
    attachment = DocumentAttachment.query.get_or_404(attachment_id)
    document = TechnicalDocument.query.get_or_404(attachment.document_id)
    
    # Verificar permissão
    if document.restricted_access and current_user.role != 'admin':
        flash('Você não tem permissão para baixar este anexo.', 'warning')
        return redirect(url_for('documents.index'))
    
    if os.path.exists(attachment.file_path):
        return send_file(
            attachment.file_path,
            as_attachment=True,
            download_name=attachment.original_filename
        )
    else:
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.view_document', document_id=document.id))


@documents_bp.route('/search', methods=['GET', 'POST'])
@login_required
def search_documents():
    """Busca avançada de documentos."""
    form = DocumentSearchForm()
    
    if form.validate_on_submit() or request.method == 'GET' and request.args.get('search_term'):
        # Obter parâmetros de busca do formulário ou da URL
        if request.method == 'POST':
            search_term = form.search_term.data
            document_type = form.document_type.data
            category = form.category.data
            status = form.status.data
            author = form.author.data
            tag = form.tag.data
        else:
            search_term = request.args.get('search_term', '')
            document_type = request.args.get('document_type', '')
            category = request.args.get('category', '')
            status = request.args.get('status', '')
            author = request.args.get('author', '')
            tag = request.args.get('tag', '')
            
            # Preencher o formulário com os valores da URL
            form.search_term.data = search_term
            form.document_type.data = document_type
            form.category.data = category
            form.status.data = status
            form.author.data = author
            form.tag.data = tag
        
        # Construir a consulta base
        query = TechnicalDocument.query
        
        # Adicionar filtros conforme os parâmetros fornecidos
        if search_term:
            search_term = f"%{search_term}%"
            query = query.filter(
                or_(
                    TechnicalDocument.title.ilike(search_term),
                    TechnicalDocument.description.ilike(search_term),
                    TechnicalDocument.original_filename.ilike(search_term)
                )
            )
        
        if document_type:
            query = query.filter(TechnicalDocument.document_type == document_type)
        
        if category:
            query = query.filter(TechnicalDocument.category == category)
        
        if status:
            query = query.filter(TechnicalDocument.status == status)
            
        if author:
            query = query.filter(TechnicalDocument.author.ilike(f"%{author}%"))
            
        if tag:
            query = query.filter(TechnicalDocument.tags.ilike(f"%{tag}%"))
        
        # Ordenar por data de upload (mais recentes primeiro)
        query = query.order_by(TechnicalDocument.upload_date.desc())
        
        # Executar a busca
        results = query.all()
        
        return render_template(
            'documents/search_results.html',
            title='Resultados da Busca',
            form=form,
            results=results,
            search_term=search_term,
            count=len(results)
        )
    
    return render_template(
        'documents/search.html',
        title='Busca Avançada de Documentos',
        form=form
    )


@documents_bp.route('/delete-attachment/<int:attachment_id>', methods=['POST'])
@login_required
def delete_attachment(attachment_id):
    """Excluir um anexo de documento."""
    attachment = DocumentAttachment.query.get_or_404(attachment_id)
    document = TechnicalDocument.query.get_or_404(attachment.document_id)
    
    # Verificar permissão
    if document.uploaded_by != current_user.id and current_user.role != 'admin':
        flash('Você não tem permissão para excluir este anexo.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Tentar excluir o arquivo físico
    if os.path.exists(attachment.file_path):
        try:
            os.remove(attachment.file_path)
        except Exception as e:
            flash(f'Erro ao excluir arquivo físico: {str(e)}', 'warning')
    
    # Excluir o registro do banco de dados
    db.session.delete(attachment)
    db.session.commit()
    
    flash('Anexo excluído com sucesso!', 'success')
    return redirect(url_for('documents.view_document', document_id=document.id))


@documents_bp.route('/change-status/<int:document_id>/<string:status>', methods=['POST'])
@login_required
def change_document_status(document_id, status):
    """Alterar o status de um documento."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão
    if document.uploaded_by != current_user.id and current_user.role != 'admin':
        flash('Você não tem permissão para alterar o status deste documento.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Verificar status válido
    if status not in ['ativo', 'em_revisao', 'obsoleto']:
        flash('Status inválido.', 'danger')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Atualizar o status
    document.status = status
    db.session.commit()
    
    flash('Status do documento atualizado com sucesso!', 'success')
    return redirect(url_for('documents.view_document', document_id=document.id))


@documents_bp.route('/create-online', methods=['GET', 'POST'])
@login_required
def create_document_online():
    """Criar documento online sem a necessidade de upload de arquivo."""
    form = DocumentForm()
    
    # Não exigir arquivo para o documento online
    form.document_file.validators = []
    
    if form.validate_on_submit():
        upload_folder = current_app.config['UPLOAD_FOLDER']
        document_folder = os.path.join(upload_folder, 'documents')
        os.makedirs(document_folder, exist_ok=True)
        
        # Criar arquivo HTML vazio para o documento online
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_documento_online.html"
        file_path = os.path.join(document_folder, filename)
        
        # Conteúdo inicial do documento HTML
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{form.title.data}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #0066cc; }}
        .container {{ max-width: 800px; margin: 0 auto; }}
        .footer {{ margin-top: 30px; font-size: 12px; color: #666; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{form.title.data}</h1>
        <p>{form.description.data if form.description.data else 'Documento criado em ' + datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        <!-- Conteúdo do documento -->
        <div class="content">
            <p>Edite este documento para adicionar seu conteúdo...</p>
        </div>
        <div class="footer">
            <p>Documento criado por: {current_user.name} - {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        </div>
    </div>
</body>
</html>
"""
        
        # Salvar o arquivo HTML
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        file_size = os.path.getsize(file_path)
        
        # Criar registro no banco de dados
        document = TechnicalDocument(
            title=form.title.data,
            description=form.description.data,
            document_type=form.document_type.data,
            category=form.category.data if form.category.data else None,
            revision=form.revision.data,
            valid_until=form.valid_until.data,
            author=form.author.data,
            tags=form.tags.data,
            status=form.status.data,
            restricted_access=form.restricted_access.data,
            filename=filename,
            original_filename=f"{form.title.data}.html",
            file_path=file_path,
            file_type='html',
            file_size=file_size,
            uploaded_by=current_user.id,
            version=1
        )
        
        db.session.add(document)
        db.session.commit()
        
        flash('Documento online criado com sucesso!', 'success')
        return redirect(url_for('documents.edit_online', document_id=document.id))
    
    return render_template(
        'documents/create_online.html',
        title='Criar Documento Online',
        form=form
    )


@documents_bp.route('/edit-online/<int:document_id>', methods=['GET', 'POST'])
@login_required
def edit_online(document_id):
    """Editar documento online."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão
    if document.uploaded_by != current_user.id and current_user.role != 'admin':
        flash('Você não tem permissão para editar este documento.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Verificar se é um documento HTML
    if document.file_type != 'html':
        flash('Este documento não pode ser editado online.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Ler o conteúdo atual do documento
    try:
        with open(document.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        flash(f'Erro ao ler o documento: {str(e)}', 'danger')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    if request.method == 'POST':
        new_content = request.form.get('content', '')
        
        # Salvar o novo conteúdo
        try:
            with open(document.file_path, 'w', encoding='utf-8') as f:
                f.write(new_content)
            
            # Atualizar tamanho do arquivo
            document.file_size = os.path.getsize(document.file_path)
            document.updated_at = datetime.datetime.utcnow()
            db.session.commit()
            
            flash('Documento atualizado com sucesso!', 'success')
        except Exception as e:
            flash(f'Erro ao salvar o documento: {str(e)}', 'danger')
        
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    return render_template(
        'documents/edit_online.html',
        title=f'Editor Online: {document.title}',
        document=document,
        content=content
    )


@documents_bp.route('/advanced-editor/<int:document_id>', methods=['GET', 'POST'])
@login_required
def advanced_editor(document_id):
    """Editor avançado de documentos com recursos modernos."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão
    if document.uploaded_by != current_user.id and current_user.role != 'admin':
        flash('Você não tem permissão para editar este documento.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Verificar se é um documento HTML
    if document.file_type != 'html':
        flash('Este documento não pode ser editado com o editor avançado.', 'warning')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    # Ler o conteúdo atual do documento
    try:
        with open(document.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
            # Extrair apenas o conteúdo do corpo se for um HTML completo
            if '<body' in content and '</body>' in content:
                body_start = content.find('<body')
                body_start = content.find('>', body_start) + 1
                body_end = content.find('</body>', body_start)
                if body_start > 0 and body_end > 0:
                    content = content[body_start:body_end].strip()
    except Exception as e:
        current_app.logger.error(f"Erro ao ler documento: {str(e)}")
        flash(f'Erro ao ler o documento: {str(e)}', 'danger')
        return redirect(url_for('documents.view_document', document_id=document.id))
    
    if request.method == 'POST':
        new_content = request.form.get('content', '')
        status = request.form.get('status', 'draft')
        
        try:
            # Envolver em um template HTML completo
            html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{document.title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #0066cc; }}
        .container {{ max-width: 100%; margin: 0 auto; }}
        .footer {{ margin-top: 30px; font-size: 12px; color: #666; }}
    </style>
</head>
<body>
    <div class="container">
        {new_content}
        <div class="footer">
            <p>Documento atualizado por: {current_user.name} - {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        </div>
    </div>
</body>
</html>"""
            
            # Salvar o conteúdo HTML completo
            with open(document.file_path, 'w', encoding='utf-8') as f:
                f.write(html_template)
            
            # Atualizar metadados do documento
            document.file_size = os.path.getsize(document.file_path)
            document.updated_at = datetime.datetime.utcnow()
            
            if status == 'published':
                document.status = 'published'
            
            db.session.commit()
            
            # Retornar resposta JSON se solicitado
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'success': True,
                    'message': 'Documento atualizado com sucesso',
                    'redirect_url': url_for('documents.view_document', document_id=document.id)
                })
            
            flash('Documento atualizado com sucesso!', 'success')
            return redirect(url_for('documents.view_document', document_id=document.id))
        
        except Exception as e:
            current_app.logger.error(f"Erro ao salvar documento: {str(e)}")
            
            # Retornar resposta JSON se solicitado
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': f'Erro ao salvar o documento: {str(e)}'})
            
            flash(f'Erro ao salvar o documento: {str(e)}', 'danger')
    
    # Gerar token CSRF para requisições Ajax
    csrf_token = generate_csrf()
    
    return render_template(
        'documents/advanced_editor.html',
        title=f'Editor Avançado: {document.title}',
        document=document,
        content=content,
        csrf_token=csrf_token
    )


# Funções para impressão avançada de documentos

@documents_bp.route('/advanced-print/<int:document_id>')
@login_required
def advanced_print_document(document_id):
    """Modo de impressão dedicado para um documento."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    if not document.file_path or not os.path.exists(document.file_path):
        flash('Arquivo não encontrado. O documento pode estar corrompido.', 'danger')
        return redirect(url_for('documents.view_document', document_id=document_id))
    
    # Verificar extensão do arquivo
    _, file_ext = os.path.splitext(document.file_path)
    file_ext = file_ext.lower()
    
    # Configuração para impressão
    print_mode = True
    auto_print = request.args.get('auto_print', 'false') == 'true'
    paper_size = request.args.get('paper_size', 'a4')
    print_header = request.args.get('header', 'true') == 'true'
    print_footer = request.args.get('footer', 'true') == 'true'
    
    # Para arquivos PDF, exibir usando um visualizador embutido
    if file_ext == '.pdf':
        return render_template(
            'documents/print_view.html',
            title=document.title,
            document=document,
            pdf_path=url_for('documents.serve_document', document_id=document_id, direct_view=True),
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )
    
    # Para imagens, mostrar diretamente
    elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
        return render_template(
            'documents/print_view.html',
            title=document.title,
            document=document,
            image_path=url_for('documents.serve_document', document_id=document_id, direct_view=True),
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )
    
    # Para arquivos Excel, tentar embutir visualizador do Microsoft Office Online
    elif file_ext in ['.xlsx', '.xls']:
        file_url = url_for('documents.serve_document', document_id=document_id, direct_view=True, _external=True)
        office_online_viewer = f"https://view.officeapps.live.com/op/embed.aspx?src={file_url}"
        
        return render_template(
            'documents/print_view.html',
            title=document.title,
            document=document,
            office_url=office_online_viewer,
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )
    
    # Para outros tipos, oferecer link de download
    else:
        download_link = url_for('documents.download_document', document_id=document_id)
        return render_template(
            'documents/print_view.html',
            title=document.title,
            document=document,
            error_message="Este tipo de arquivo não pode ser impresso diretamente.",
            download_link=download_link,
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )


@documents_bp.route('/advanced-print-virtual/<path:file_path>')
@login_required
def advanced_print_virtual_document(file_path):
    """Modo de impressão dedicado para um documento virtual - mostra apenas o documento."""
    # Montar o caminho completo
    FORMS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'extracted_forms')
    
    # Garantir que o caminho é seguro (dentro do diretório permitido)
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('Arquivo não encontrado.', 'danger')
        return redirect(url_for('documents.index'))
        
    # Criar informações do documento
    file_name = os.path.basename(file_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    # Informações do documento (simplificado para impressão)
    doc_info = {
        'title': file_name,
        'file_path': full_path,
        'created_at': datetime.datetime.fromtimestamp(os.path.getctime(full_path)),
        'file_type': file_ext[1:] if file_ext.startswith('.') else file_ext,
        'author': 'Sistema Zelopack',
        'category': 'Documento'
    }
    
    # Link para download
    download_link = url_for('documents.download_virtual_document', file_path=file_path)
    
    # Configuração para impressão
    print_mode = True
    auto_print = request.args.get('auto_print', 'false') == 'true'
    paper_size = request.args.get('paper_size', 'a4')
    print_header = request.args.get('header', 'true') == 'true'
    print_footer = request.args.get('footer', 'true') == 'true'
    
    # Determinar como mostrar com base na extensão do arquivo
    if file_ext in ['.pdf']:
        file_url = url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false')
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            pdf_path=file_url,
            download_link=download_link,
            virtual_document=True,
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )
    elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
        file_url = url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false')
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            image_path=file_url,
            download_link=download_link,
            virtual_document=True,
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )
    elif file_ext in ['.xlsx', '.xls']:
        file_url = url_for('documents.download_virtual_document', file_path=file_path, as_attachment='false', _external=True)
        office_viewer_url = f"https://view.officeapps.live.com/op/embed.aspx?src={file_url}"
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            office_url=office_viewer_url,
            download_link=download_link,
            virtual_document=True,
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )
    else:
        return render_template(
            'documents/print_view.html',
            title=f'Impressão: {file_name}',
            document=doc_info,
            error_message="Este tipo de arquivo não pode ser impresso diretamente.",
            download_link=download_link,
            virtual_document=True,
            print_mode=print_mode,
            auto_print=auto_print,
            paper_size=paper_size,
            print_header=print_header,
            print_footer=print_footer
        )


@documents_bp.route('/image-preview/<int:document_id>')
@login_required
def image_preview(document_id):
    """Gerar uma prévia da imagem para documentos do tipo imagem ou PDF."""
    document = TechnicalDocument.query.get_or_404(document_id)
    
    # Verificar permissão para documentos restritos
    if document.restricted_access and current_user.role != 'admin':
        abort(403)
    
    # Constante para tamanho da thumbnail
    THUMBNAIL_SIZE = (400, 400)
    
    # Verificar se o documento é uma imagem
    image_extensions = ['jpg', 'jpeg', 'png', 'gif', 'bmp']
    
    try:
        if document.file_type.lower() in image_extensions:
            # Processar imagens diretas
            img = Image.open(document.file_path)
            img.thumbnail(THUMBNAIL_SIZE)
            
            img_io = io.BytesIO()
            img.save(img_io, format=img.format or 'JPEG')
            img_io.seek(0)
            
            return send_file(img_io, mimetype=f'image/{img.format.lower() if img.format else "jpeg"}')
        
        elif document.file_type.lower() == 'pdf':
            # Processar PDFs - gerar uma imagem da primeira página
            try:
                from pdf2image import convert_from_path
                import tempfile
                
                # Gerar imagem da primeira página apenas
                with tempfile.TemporaryDirectory() as path:
                    images = convert_from_path(document.file_path, dpi=72, output_folder=path, 
                                               first_page=1, last_page=1, size=THUMBNAIL_SIZE)
                    
                    if images and len(images) > 0:
                        img = images[0]
                        img_io = io.BytesIO()
                        img.save(img_io, format='JPEG')
                        img_io.seek(0)
                        
                        return send_file(img_io, mimetype='image/jpeg')
            except ImportError:
                # Se pdf2image não estiver disponível, tente usar PyPDF2 + reportlab para renderizar
                try:
                    from PyPDF2 import PdfReader
                    from reportlab.lib.pagesizes import letter
                    from reportlab.pdfgen import canvas
                    
                    pdf_reader = PdfReader(document.file_path)
                    if len(pdf_reader.pages) > 0:
                        # Criar imagem com texto "Prévia do PDF" no centro
                        img = Image.new('RGB', (300, 400), color=(245, 245, 245))
                        draw = ImageDraw.Draw(img)
                        
                        # Tente carregar uma fonte, ou use a fonte padrão
                        try:
                            font = ImageFont.truetype("arial.ttf", 16)
                        except IOError:
                            font = ImageFont.load_default()
                        
                        # Adicione o título do documento
                        title = document.title
                        if len(title) > 25:
                            title = title[:22] + "..."
                        
                        # Desenhe informações do PDF
                        draw.text((20, 20), f"PDF: {title}", fill=(0, 0, 0), font=font)
                        draw.text((20, 50), f"Páginas: {len(pdf_reader.pages)}", fill=(0, 0, 0), font=font)
                        draw.text((20, 80), "Prévia do documento", fill=(0, 0, 0), font=font)
                        
                        # Adicione um ícone de PDF (retângulo vermelho com "PDF")
                        draw.rectangle((100, 150, 200, 250), fill=(220, 50, 50))
                        draw.text((130, 190), "PDF", fill=(255, 255, 255), font=font)
                        
                        # Converta a imagem para enviar
                        img_io = io.BytesIO()
                        img.save(img_io, format='JPEG')
                        img_io.seek(0)
                        
                        return send_file(img_io, mimetype='image/jpeg')
                except (ImportError, Exception) as e:
                    # Se tudo falhar, retorne um ícone de PDF padrão
                    current_app.logger.error(f"Erro ao gerar prévia de PDF: {str(e)}")
            
        # Para qualquer outro tipo de arquivo ou se as conversões falharem
        # Retornar um ícone padrão com base no tipo
        img = Image.new('RGB', (300, 300), color=(245, 245, 245))
        draw = ImageDraw.Draw(img)
        
        # Tente carregar uma fonte
        try:
            font = ImageFont.truetype("arial.ttf", 16)
        except IOError:
            font = ImageFont.load_default()
        
        # Informações básicas
        if document.file_type.lower() == 'pdf':
            draw.text((20, 20), "Documento PDF", fill=(0, 0, 0), font=font)
            color = (220, 50, 50)  # Vermelho para PDF
        elif document.file_type.lower() in ['doc', 'docx']:
            draw.text((20, 20), "Documento Word", fill=(0, 0, 0), font=font)
            color = (50, 50, 220)  # Azul para Word
        elif document.file_type.lower() in ['xls', 'xlsx']:
            draw.text((20, 20), "Planilha Excel", fill=(0, 0, 0), font=font)
            color = (50, 150, 50)  # Verde para Excel
        else:
            draw.text((20, 20), f"Arquivo {document.file_type.upper()}", fill=(0, 0, 0), font=font)
            color = (150, 150, 150)  # Cinza para outros tipos
        
        # Título do documento
        title = document.title
        if len(title) > 25:
            title = title[:22] + "..."
        draw.text((20, 50), title, fill=(0, 0, 0), font=font)
        
        # Ícone representativo
        draw.rectangle((100, 100, 200, 200), fill=color)
        draw.text((130, 140), document.file_type.upper(), fill=(255, 255, 255), font=font)
        
        # Converta e envie a imagem
        img_io = io.BytesIO()
        img.save(img_io, format='JPEG')
        img_io.seek(0)
        
        return send_file(img_io, mimetype='image/jpeg')
    
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar prévia da imagem: {str(e)}")
        abort(500)