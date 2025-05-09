"""
Módulo para processamento de diferentes tipos de documentos/formulários.
"""
import os
import io
import re
import openpyxl
from docx import Document
import PyPDF2
import pandas as pd

def process_xlsx(file_path):
    """
    Processa uma planilha Excel e retorna uma estrutura de dados para edição.
    """
    if not os.path.exists(file_path) or not file_path.lower().endswith('.xlsx'):
        raise ValueError("Arquivo inválido ou não é uma planilha Excel.")
    
    result = {
        'type': 'xlsx',
        'filename': os.path.basename(file_path),
        'sheets': [],
    }
    
    # Carregar a planilha
    wb = openpyxl.load_workbook(file_path, data_only=True)
    
    # Processar cada planilha
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_data = {
            'name': sheet_name,
            'rows': sheet.max_row,
            'columns': sheet.max_column,
            'cells': [],
            'tables': [],  # Para tabelas identificadas
            'headers': [],  # Para cabeçalhos identificados
            'fields': [],   # Para campos identificados
        }
        
        # Identificar cabeçalhos e rodapés
        for row in range(1, min(10, sheet.max_row + 1)):  # Procurar nos primeiros 10 rows
            for col in range(1, min(10, sheet.max_column + 1)):
                cell = sheet.cell(row=row, column=col)
                if cell.value and isinstance(cell.value, str) and cell.font.bold:
                    sheet_data['headers'].append({
                        'row': row,
                        'column': col,
                        'value': str(cell.value),
                    })
        
        # Identificar possíveis campos de formulário
        for row in range(1, sheet.max_row + 1):
            for col in range(1, sheet.max_column + 1):
                cell = sheet.cell(row=row, column=col)
                if cell.value and isinstance(cell.value, str):
                    text = str(cell.value).strip()
                    next_cell = None
                    
                    # Verificar se a célula parece um rótulo de campo
                    if (
                        text.endswith(':') or 
                        re.search(r'(Data|Nome|Valor|Quantidade|Observação|Responsável)$', text, re.IGNORECASE)
                    ):
                        # Verificar a célula à direita ou abaixo para entrada de dados
                        if col < sheet.max_column:
                            next_cell = sheet.cell(row=row, column=col+1)
                        elif row < sheet.max_row:
                            next_cell = sheet.cell(row=row+1, column=col)
                        
                        field_id = f"{sheet_name}_field_{row}_{col}"
                        field_type = "text"  # Tipo padrão
                        
                        # Tentar determinar o tipo de campo
                        if text.lower().find('data') >= 0:
                            field_type = "date"
                        elif text.lower().find('valor') >= 0 or text.lower().find('quantidade') >= 0:
                            field_type = "number"
                        elif text.lower().find('observação') >= 0 or text.lower().find('descrição') >= 0:
                            field_type = "textarea"
                        elif text.lower().find('sim/não') >= 0 or text.lower().find('aprovado') >= 0:
                            field_type = "boolean"
                        
                        field = {
                            'id': field_id,
                            'label': text,
                            'type': field_type,
                            'row': row,
                            'column': col,
                            'value': next_cell.value if next_cell and next_cell.value else "",
                            'editable': True,
                        }
                        
                        sheet_data['fields'].append(field)
        
        # Adicionar dados da planilha ao resultado
        result['sheets'].append(sheet_data)
    
    return result

def process_docx(file_path):
    """
    Processa um documento Word e retorna uma estrutura de dados para edição.
    """
    if not os.path.exists(file_path) or not file_path.lower().endswith('.docx'):
        raise ValueError("Arquivo inválido ou não é um documento Word.")
    
    result = {
        'type': 'docx',
        'filename': os.path.basename(file_path),
        'paragraphs': [],
        'tables': [],
        'fields': [],
    }
    
    # Carregar o documento
    doc = Document(file_path)
    
    # Processar parágrafos
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # Ignorar parágrafos vazios
            # Verificar se o parágrafo tem estilo de título
            is_heading = para.style.name.startswith('Heading')
            
            # Verificar se parece ser um campo de formulário
            is_field = False
            field_id = None
            field_type = "text"
            
            if ':' in text:
                parts = text.split(':', 1)
                label = parts[0].strip()
                value = parts[1].strip() if len(parts) > 1 else ""
                
                # Criar um ID único para o campo
                field_id = f"para_{i}_{label.lower().replace(' ', '_')}"
                is_field = True
                
                # Tentar determinar o tipo de campo
                if label.lower().find('data') >= 0:
                    field_type = "date"
                elif label.lower().find('valor') >= 0 or label.lower().find('quantidade') >= 0:
                    field_type = "number"
                elif label.lower().find('observação') >= 0 or label.lower().find('descrição') >= 0:
                    field_type = "textarea"
                
                # Adicionar campo à lista de campos identificados
                if is_field:
                    result['fields'].append({
                        'id': field_id,
                        'label': label,
                        'value': value,
                        'type': field_type,
                        'paragraph': i,
                        'editable': True,
                    })
            
            # Adicionar parágrafo aos dados de resultado
            result['paragraphs'].append({
                'index': i,
                'text': text,
                'is_heading': is_heading,
                'is_field': is_field,
                'field_id': field_id if is_field else None,
            })
    
    # Processar tabelas
    for t_idx, table in enumerate(doc.tables):
        table_data = {
            'index': t_idx,
            'rows': len(table.rows),
            'columns': len(table.rows[0].cells) if table.rows else 0,
            'cells': [],
            'fields': [],
        }
        
        # Processar células
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                
                # Verificar se a célula parece ser um rótulo de campo
                is_field = False
                field_id = None
                
                if text.endswith(':') or re.search(r'(Data|Nome|Valor|Quantidade|Observação)$', text, re.IGNORECASE):
                    # Se a próxima célula existe e está vazia, é um candidato a campo
                    next_cell = row.cells[c_idx + 1] if c_idx + 1 < len(row.cells) else None
                    
                    if next_cell and not next_cell.text.strip():
                        field_id = f"table_{t_idx}_row_{r_idx}_col_{c_idx}"
                        is_field = True
                        field_type = "text"  # Tipo padrão
                        
                        # Tentar determinar o tipo de campo
                        if text.lower().find('data') >= 0:
                            field_type = "date"
                        elif text.lower().find('valor') >= 0 or text.lower().find('quantidade') >= 0:
                            field_type = "number"
                        
                        table_data['fields'].append({
                            'id': field_id,
                            'label': text,
                            'value': "",
                            'type': field_type,
                            'row': r_idx,
                            'column': c_idx,
                            'editable': True,
                        })
                
                # Adicionar célula aos dados da tabela
                table_data['cells'].append({
                    'row': r_idx,
                    'column': c_idx,
                    'text': text,
                    'is_field': is_field,
                    'field_id': field_id if is_field else None,
                })
        
        # Adicionar tabela ao resultado
        result['tables'].append(table_data)
    
    return result

def process_pdf(file_path):
    """
    Processa um arquivo PDF e retorna uma estrutura de dados para edição.
    """
    if not os.path.exists(file_path) or not file_path.lower().endswith('.pdf'):
        raise ValueError("Arquivo inválido ou não é um PDF.")
    
    result = {
        'type': 'pdf',
        'filename': os.path.basename(file_path),
        'pages': [],
        'fields': [],
        'interactive': False,
    }
    
    # Abrir o PDF
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        
        # Verificar se é um PDF interativo (tem campos de formulário)
        pdf_fields = reader.get_fields()
        result['interactive'] = bool(pdf_fields)
        
        if pdf_fields:
            # Processar campos interativos
            for field_name, field in pdf_fields.items():
                field_type = "text"  # Tipo padrão
                
                # Determinar o tipo de campo com base no nome ou propriedades
                if field_name.lower().find('data') >= 0:
                    field_type = "date"
                elif field_name.lower().find('valor') >= 0 or field_name.lower().find('quantidade') >= 0:
                    field_type = "number"
                elif field_name.lower().find('check') >= 0 or field_name.lower().find('sim_nao') >= 0:
                    field_type = "checkbox"
                
                result['fields'].append({
                    'id': field_name,
                    'name': field_name,
                    'type': field_type,
                    'value': str(field.get('/V', '')),
                    'editable': True,
                })
        
        # Processar páginas do PDF
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            lines = text.split('\n')
            
            # Processar linhas para identificar possíveis campos
            fields_in_page = []
            for line_num, line in enumerate(lines):
                if ':' in line:
                    parts = line.split(':', 1)
                    label = parts[0].strip()
                    value = parts[1].strip() if len(parts) > 1 else ""
                    
                    # Criar campo apenas se o label não for muito grande
                    if len(label) < 50:
                        field_id = f"page_{page_num+1}_line_{line_num+1}"
                        field_type = "text"  # Tipo padrão
                        
                        # Tentar determinar o tipo de campo
                        if label.lower().find('data') >= 0:
                            field_type = "date"
                        elif label.lower().find('valor') >= 0 or label.lower().find('quantidade') >= 0:
                            field_type = "number"
                        
                        field = {
                            'id': field_id,
                            'label': label,
                            'value': value,
                            'type': field_type,
                            'page': page_num + 1,
                            'line': line_num + 1,
                            'editable': not result['interactive'],  # Editável apenas se não for interativo
                        }
                        
                        fields_in_page.append(field)
                        # Adicionar também à lista global de campos
                        if not result['interactive']:
                            result['fields'].append(field)
            
            # Adicionar dados da página ao resultado
            result['pages'].append({
                'number': page_num + 1,
                'text': text,
                'lines': lines,
                'fields': fields_in_page,
            })
    
    return result