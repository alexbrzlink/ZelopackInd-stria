"""
Módulo para extração de campos e estrutura de documentos e formulários.
"""
import os
import re
import openpyxl
from docx import Document
import PyPDF2

def extract_form_fields(file_path):
    """
    Extrai campos de formulários baseado no tipo de arquivo.
    Retorna uma lista de campos encontrados.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.xlsx':
        return extract_xlsx_fields(file_path)
    elif file_extension == '.docx':
        return extract_docx_fields(file_path)
    elif file_extension == '.pdf':
        return extract_pdf_fields(file_path)
    else:
        raise ValueError(f"Tipo de arquivo não suportado: {file_extension}")

def extract_xlsx_fields(file_path):
    """Extrai campos de um arquivo Excel."""
    fields = []
    wb = openpyxl.load_workbook(file_path, data_only=True)
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        # Procurar por células que possam ser campos de formulário
        for row in range(1, sheet.max_row + 1):
            for col in range(1, sheet.max_column + 1):
                cell = sheet.cell(row=row, column=col)
                
                # Verificar se o valor da célula tem indicativos de ser um campo
                if cell.value and isinstance(cell.value, str):
                    text = cell.value.strip()
                    
                    # Padrões comuns para campos de formulário
                    if (
                        ':' in text or 
                        re.search(r'\(.*\)', text) or  # Texto entre parênteses
                        text.endswith('?') or  # Perguntas
                        text.lower().startswith(('nome', 'data', 'valor', 'quantidade', 'observações'))
                    ):
                        fields.append({
                            'sheet': sheet_name,
                            'position': f"{col},{row}",
                            'text': text,
                            'type': 'input',  # Tipo padrão, poderia ser refinado
                            'value': '',  # Valor vazio para preenchimento
                        })
    
    return fields

def extract_docx_fields(file_path):
    """Extrai campos de um documento Word."""
    fields = []
    doc = Document(file_path)
    
    # Processar parágrafos
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Identificar possíveis campos
        if text and (':' in text or text.endswith('?')):
            fields.append({
                'paragraph': i,
                'text': text,
                'type': 'input',
                'value': '',
            })
    
    # Processar tabelas
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                
                # Identificar possíveis campos em células de tabela
                if text and (':' in text or re.search(r'\(.*\)', text) or text.endswith('?')):
                    fields.append({
                        'table': t_idx,
                        'row': r_idx,
                        'col': c_idx,
                        'text': text,
                        'type': 'input',
                        'value': '',
                    })
    
    return fields

def extract_pdf_fields(file_path):
    """Extrai campos de um arquivo PDF."""
    fields = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        
        # Verificar se o PDF tem campos de formulário interativos
        if reader.get_fields():
            for field_name, field_value in reader.get_fields().items():
                fields.append({
                    'name': field_name,
                    'type': 'pdf_field',
                    'value': str(field_value),
                })
        else:
            # PDF não interativo, tentar extrair texto e identificar campos
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()
                
                # Quebrar em linhas
                lines = text.split('\n')
                for line_num, line in enumerate(lines):
                    # Identificar possíveis campos em linhas de texto
                    if ':' in line or re.search(r'\(.*\)', line) or line.endswith('?'):
                        fields.append({
                            'page': page_num + 1,
                            'line': line_num + 1,
                            'text': line.strip(),
                            'type': 'input',
                            'value': '',
                        })
    
    return fields

def extract_form_structure(file_path):
    """
    Analisa a estrutura geral do documento e retorna uma representação.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Informações básicas do arquivo
    stats = os.stat(file_path)
    structure = {
        'filename': os.path.basename(file_path),
        'size': stats.st_size,
        'modified': stats.st_mtime,
        'type': file_extension[1:],  # Remove o ponto inicial
    }
    
    # Extração específica por tipo
    if file_extension == '.xlsx':
        wb = openpyxl.load_workbook(file_path, data_only=True)
        structure['sheets'] = wb.sheetnames
        structure['cell_count'] = sum(sheet.max_row * sheet.max_column for sheet in wb.worksheets)
        
    elif file_extension == '.docx':
        doc = Document(file_path)
        structure['paragraphs'] = len(doc.paragraphs)
        structure['tables'] = len(doc.tables)
        
    elif file_extension == '.pdf':
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            structure['pages'] = len(reader.pages)
            structure['has_form_fields'] = bool(reader.get_fields())
    
    return structure